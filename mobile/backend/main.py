import os
import io
import re
import json
import uuid
import base64
import asyncio
import logging
import requests as http_requests
from datetime import datetime, timedelta, date
from pathlib import Path
from fastapi import FastAPI, HTTPException, BackgroundTasks, UploadFile, File, Depends, Header, Query, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel, Field
from openai import OpenAI
from pypdf import PdfReader
from docx import Document as DocxDocument
from sqlalchemy import text
from sqlalchemy.orm import Session
from tavily import TavilyClient

# Import database Base for Alembic to discover models
from database import Base, SessionLocal, get_db, engine
import models  # Import models so Alembic can see them
from auth import hash_password, decode_token, verify_password, create_access_token
from microsoft_auth import verify_microsoft_token
from rag import index_document, get_relevant_context, get_top_document, get_top_chunks_for_document, get_chunks_by_filename, list_indexed_filenames, warm_reranker
from websocket_manager import manager
from blob_storage import is_blob_storage_enabled, upload_blob, download_blob, delete_blob

app = FastAPI(title="Group Chat Prototype")

# CORS configuration - allow localhost for development and frontend URL from environment for production
frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
allowed_origins = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
    "https://ms3015-chatbot.vercel.app"
]

# Add frontend URL from environment if it's different from localhost
if frontend_url not in allowed_origins:
    allowed_origins.append(frontend_url)

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


EXPO_PUSH_URL = "https://exp.host/--/api/v2/push/send"


def send_push_notifications(tokens: list[str], title: str, body: str, data: dict):
    """Fire-and-forget push to the Expo push service. Failures are non-critical."""
    messages = [
        {"to": t, "title": title, "body": body, "data": data, "sound": "default"}
        for t in tokens
        if t.startswith("ExponentPushToken[")
    ]
    if not messages:
        return
    try:
        http_requests.post(EXPO_PUSH_URL, json=messages, timeout=5)
    except Exception:
        pass  # push failures should never break the WS flow


def init_demo_data():
    """
    Initialize demo data in the database if it's empty.
    Only runs if no users exist in the database.
    """
    db = SessionLocal()
    try:
        # Check if any users exist
        existing_user = db.query(models.User).first()
        if existing_user:
            print("Database already has data. Skipping demo data initialization.")
            return

        print("Initializing demo data...")

        # Create Users
        # 1 coordinator
        coordinator = models.User(
            username="coordinator",
            password_hash=hash_password("coordinator1"),
            role="coordinator"
        )
        db.add(coordinator)

        # 2 supervisors
        supervisor1 = models.User(
            username="supervisor1",
            password_hash=hash_password("supervisor1"),
            role="supervisor"
        )
        supervisor2 = models.User(
            username="supervisor2",
            password_hash=hash_password("supervisor2"),
            role="supervisor"
        )
        db.add(supervisor1)
        db.add(supervisor2)

        # 8 students
        students = []
        for i in range(1, 9):
            student = models.User(
                username=f"student{i}",
                password_hash=hash_password(f"student{i}"),
                role="student"
            )
            students.append(student)
            db.add(student)

        # Flush to get IDs assigned
        db.flush()

        # Create Groups with string_id mapping
        group1 = models.Group(name="Group 1", string_id="group-a")
        group2 = models.Group(name="Group 2", string_id="group-b")
        group3 = models.Group(name="Group 3", string_id="group-c")
        group4 = models.Group(name="Group 4", string_id="group-d")
        db.add_all([group1, group2, group3, group4])
        db.flush()
        
        # Seed a few initial messages (only if no messages exist)
        if db.query(models.Message).count() == 0:
            # Group 1 messages
            db.add(models.Message(
                group_id=group1.id,
                user_id=students[0].id,
                content="Hey everyone!",
                is_AI=False
            ))
            db.add(models.Message(
                group_id=group1.id,
                user_id=students[1].id,
                content="Hi Alice!",
                is_AI=False
            ))
            # Group 2 message
            db.add(models.Message(
                group_id=group2.id,
                user_id=students[2].id,
                content="Who's ready for lunch?",
                is_AI=False
            ))
            db.flush()

        # Create Group Memberships
        # Group 1: supervisor1, student1, student2
        db.add(models.GroupMember(user_id=supervisor1.id, group_id=group1.id))
        db.add(models.GroupMember(user_id=students[0].id, group_id=group1.id))
        db.add(models.GroupMember(user_id=students[1].id, group_id=group1.id))

        # Group 2: supervisor1, student3, student4
        db.add(models.GroupMember(user_id=supervisor1.id, group_id=group2.id))
        db.add(models.GroupMember(user_id=students[2].id, group_id=group2.id))
        db.add(models.GroupMember(user_id=students[3].id, group_id=group2.id))

        # Group 3: supervisor2, student5, student6
        db.add(models.GroupMember(user_id=supervisor2.id, group_id=group3.id))
        db.add(models.GroupMember(user_id=students[4].id, group_id=group3.id))
        db.add(models.GroupMember(user_id=students[5].id, group_id=group3.id))

        # Group 4: supervisor2, student7, student8
        db.add(models.GroupMember(user_id=supervisor2.id, group_id=group4.id))
        db.add(models.GroupMember(user_id=students[6].id, group_id=group4.id))
        db.add(models.GroupMember(user_id=students[7].id, group_id=group4.id))

        # Commit all changes
        db.commit()
        print("Demo data initialized successfully!")

    except Exception as e:
        db.rollback()
        print(f"Error initializing demo data: {str(e)}")
        raise
    finally:
        db.close()


def migrate_student_summaries():
    # runs once on startup — copies any existing Group.student_summary text into
    # the new student_summaries table so we don't lose data when switching to the new model
    db = SessionLocal()
    try:
        # already migrated on a previous startup — nothing to do
        if db.query(models.StudentSummary).first():
            return

        groups_with_text = db.query(models.Group).filter(
            models.Group.student_summary != None,
            models.Group.student_summary != ""
        ).all()

        for group in groups_with_text:
            db.add(models.StudentSummary(
                group_id=group.string_id,
                summary_text=group.student_summary,
                created_at=datetime.utcnow(),
                created_by_user_id=None,
            ))

        db.commit()
    except Exception as e:
        db.rollback()
        print(f"Error during student summary migration: {str(e)}")
    finally:
        db.close()


@app.on_event("startup")
async def startup_event():
    """Run initialization tasks when the app starts."""
    # pgvector extension must exist before create_all() tries to create the chunks table
    if engine.dialect.name == "postgresql":
        with engine.connect() as conn:
            conn.execute(text("CREATE EXTENSION IF NOT EXISTS vector"))
            conn.commit()
    # Create all database tables
    Base.metadata.create_all(bind=engine)
    # Add any columns that were added after the DB was first created.
    # SQLite doesn't support IF NOT EXISTS on ALTER TABLE, so wrap in try/except.
    with engine.connect() as conn:
        is_postgres = engine.dialect.name == "postgresql"
        if is_postgres:
            conn.execute(text("ALTER TABLE group_members ADD COLUMN IF NOT EXISTS last_read_message_id INTEGER"))
            conn.commit()
            conn.execute(text("ALTER TABLE messages ADD COLUMN IF NOT EXISTS message_type VARCHAR NOT NULL DEFAULT 'text'"))
            conn.commit()
            conn.execute(text("ALTER TABLE student_summaries ADD COLUMN IF NOT EXISTS ai_summary_copy TEXT"))
            conn.commit()
            conn.execute(text("ALTER TABLE student_summaries ADD COLUMN IF NOT EXISTS is_submitted BOOLEAN NOT NULL DEFAULT false"))
            conn.commit()
            conn.execute(text("ALTER TABLE student_summaries ADD COLUMN IF NOT EXISTS submitted_at TIMESTAMP"))
            conn.commit()
            conn.execute(text("ALTER TABLE student_summaries ADD COLUMN IF NOT EXISTS is_late BOOLEAN NOT NULL DEFAULT false"))
            conn.commit()
            conn.execute(text("ALTER TABLE deadlines ADD COLUMN IF NOT EXISTS start_dt TIMESTAMP"))
            conn.commit()
            conn.execute(text("ALTER TABLE deadlines ADD COLUMN IF NOT EXISTS frequency VARCHAR NOT NULL DEFAULT 'once'"))
            conn.commit()
            conn.execute(text("ALTER TABLE deadlines ADD COLUMN IF NOT EXISTS is_hard BOOLEAN NOT NULL DEFAULT false"))
            conn.commit()
            conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS email TEXT UNIQUE"))
            conn.commit()
            conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS full_name TEXT"))
            conn.commit()
            conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS is_active BOOLEAN NOT NULL DEFAULT true"))
            conn.commit()
            conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS student_id TEXT"))
            conn.commit()
            # chunks table was just created by create_all() above — build the HNSW index now
            conn.execute(text("CREATE INDEX IF NOT EXISTS chunks_embedding_idx ON chunks USING hnsw (embedding vector_cosine_ops)"))
            conn.commit()
        else:
            try:
                conn.execute(text("ALTER TABLE group_members ADD COLUMN last_read_message_id INTEGER"))
                conn.commit()
            except Exception:
                pass  # column already exists — safe to ignore
            try:
                conn.execute(text("ALTER TABLE messages ADD COLUMN message_type VARCHAR DEFAULT 'text' NOT NULL"))
                conn.commit()
            except Exception:
                pass  # column already exists — safe to ignore
            try:
                conn.execute(text("ALTER TABLE student_summaries ADD COLUMN ai_summary_copy TEXT"))
                conn.commit()
            except Exception:
                pass
            try:
                conn.execute(text("ALTER TABLE student_summaries ADD COLUMN is_submitted BOOLEAN NOT NULL DEFAULT 0"))
                conn.commit()
            except Exception:
                pass
            try:
                conn.execute(text("ALTER TABLE student_summaries ADD COLUMN submitted_at DATETIME"))
                conn.commit()
            except Exception:
                pass
            try:
                conn.execute(text("ALTER TABLE student_summaries ADD COLUMN is_late BOOLEAN NOT NULL DEFAULT 0"))
                conn.commit()
            except Exception:
                pass
            try:
                conn.execute(text("ALTER TABLE deadlines ADD COLUMN start_dt DATETIME"))
                conn.commit()
            except Exception:
                pass
            try:
                conn.execute(text("ALTER TABLE deadlines ADD COLUMN frequency VARCHAR NOT NULL DEFAULT 'once'"))
                conn.commit()
            except Exception:
                pass
            try:
                conn.execute(text("ALTER TABLE deadlines ADD COLUMN is_hard BOOLEAN NOT NULL DEFAULT 0"))
                conn.commit()
            except Exception:
                pass
            try:
                conn.execute(text("ALTER TABLE users ADD COLUMN email TEXT UNIQUE"))
                conn.commit()
            except Exception:
                pass
            try:
                conn.execute(text("ALTER TABLE users ADD COLUMN full_name TEXT"))
                conn.commit()
            except Exception:
                pass
            try:
                conn.execute(text("ALTER TABLE users ADD COLUMN is_active INTEGER NOT NULL DEFAULT 1"))
                conn.commit()
            except Exception:
                pass
            try:
                conn.execute(text("ALTER TABLE users ADD COLUMN student_id TEXT"))
                conn.commit()
            except Exception:
                pass
        if is_postgres:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS deadlines (
                    id SERIAL PRIMARY KEY,
                    deadline_dt TIMESTAMP NOT NULL,
                    set_by_user_id INTEGER REFERENCES users(id),
                    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
                )
            """))
        else:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS deadlines (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    deadline_dt DATETIME NOT NULL,
                    set_by_user_id INTEGER REFERENCES users(id),
                    created_at DATETIME NOT NULL DEFAULT (datetime('now'))
                )
            """))
        conn.commit()
        if is_postgres:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS course_periods (
                    id SERIAL PRIMARY KEY,
                    start_date DATE NOT NULL,
                    end_date DATE NOT NULL,
                    set_by_user_id INTEGER REFERENCES users(id),
                    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
                )
            """))
        else:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS course_periods (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    start_date DATE NOT NULL,
                    end_date DATE NOT NULL,
                    set_by_user_id INTEGER REFERENCES users(id),
                    created_at DATETIME NOT NULL DEFAULT (datetime('now'))
                )
            """))
        conn.commit()
        if is_postgres:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS broadcasts (
                    id SERIAL PRIMARY KEY,
                    content TEXT NOT NULL,
                    sent_by_user_id INTEGER REFERENCES users(id),
                    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
                )
            """))
        else:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS broadcasts (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    content TEXT NOT NULL,
                    sent_by_user_id INTEGER REFERENCES users(id),
                    created_at DATETIME NOT NULL DEFAULT (datetime('now'))
                )
            """))
        conn.commit()
        if is_postgres:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS feedback (
                    id SERIAL PRIMARY KEY,
                    content TEXT NOT NULL,
                    feedback_type TEXT NOT NULL,
                    submitted_by_user_id INTEGER REFERENCES users(id),
                    is_resolved BOOLEAN NOT NULL DEFAULT false,
                    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP
                )
            """))
        else:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS feedback (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    content TEXT NOT NULL,
                    feedback_type TEXT NOT NULL,
                    submitted_by_user_id INTEGER REFERENCES users(id),
                    is_resolved INTEGER NOT NULL DEFAULT 0,
                    created_at DATETIME NOT NULL DEFAULT (datetime('now'))
                )
            """))
        conn.commit()
    # Copy any existing Group.student_summary text into the new student_summaries table
    migrate_student_summaries()
    # Ensure uploads directories exist
    PDF_STORAGE_DIR.mkdir(parents=True, exist_ok=True)
    IMAGE_STORAGE_DIR.mkdir(parents=True, exist_ok=True)
    # Only initialize demo data if SKIP_DEMO_DATA is not set (i.e., in local development)
    if not os.getenv("SKIP_DEMO_DATA"):
        init_demo_data()
    # kick off the reranker model download in the background — startup returns immediately,
    # model is usually ready well before the first @ai query arrives
    asyncio.get_running_loop().run_in_executor(None, warm_reranker)

# Initialize OpenAI client with API key from environment variable
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    print("WARNING: OPENAI_API_KEY environment variable not set. AI features will not work.")
    openai_client = None
else:
    openai_client = OpenAI(api_key=openai_api_key)

# Setup PDF storage directory
PDF_STORAGE_DIR = Path("uploads/pdfs")
PDF_STORAGE_DIR.mkdir(parents=True, exist_ok=True)

# Setup image storage directory for chat images
IMAGE_STORAGE_DIR = Path("uploads/images")
IMAGE_STORAGE_DIR.mkdir(parents=True, exist_ok=True)

# In-memory data store for hard-coded groups and messages
groups = {
    "group-a": {
        "id": "group-a",
        "name": "Group 1",
        "messages": [
            {"id": 1, "sender": "Alice", "text": "Hey everyone!", "is_bot": False},
            {"id": 2, "sender": "Bob", "text": "Hi Alice!", "is_bot": False},
        ],
    },
    "group-b": {
        "id": "group-b",
        "name": "Group 2",
        "messages": [
            {"id": 1, "sender": "Charlie", "text": "Who's ready for lunch?", "is_bot": False},
        ],
    },
    "group-c": {
        "id": "group-c",
        "name": "Group 3",
        "messages": [],
    },
    "group-d": {
        "id": "group-d",
        "name": "Group 4",
        "messages": [],
    },
}

# In-memory storage removed - documents are now persisted in PostgreSQL


class NewMessage(BaseModel):
    sender: str
    text: str = Field(max_length=4000)


class LoginRequest(BaseModel):
    username: str
    password: str


class StudentSummaryRequest(BaseModel):
    summary_text: str
    ai_summary_copy: str | None = None


class DeadlineRequest(BaseModel):
    start_dt: datetime
    frequency: str = "once"
    is_hard: bool = False


class CoursePeriodRequest(BaseModel):
    start_date: date
    end_date: date


class CreateUserRequest(BaseModel):
    username: str
    password: str
    email: str | None = None
    full_name: str | None = None
    role: str


class UpdateUserRequest(BaseModel):
    email: str | None = None
    full_name: str | None = None
    role: str | None = None
    is_active: bool | None = None


class BulkUserEntry(BaseModel):
    username: str
    password: str
    role: str
    email: str | None = None
    full_name: str | None = None
    group_id: int | None = None


class AddGroupMemberRequest(BaseModel):
    user_id: int
    role_in_group: str = "member"


class CreateGroupRequest(BaseModel):
    name: str


class ExcelStudentEntry(BaseModel):
    full_name: str
    username: str
    student_id: str
    group_id: str


class ChangePasswordRequest(BaseModel):
    current_password: str
    new_password: str


class BroadcastRequest(BaseModel):
    content: str


class FeedbackRequest(BaseModel):
    content: str
    feedback_type: str


class MicrosoftLoginRequest(BaseModel):
    id_token: str


def compute_next_deadline(deadline_row) -> datetime:
    # fall back to deadline_dt for rows created before start_dt existed
    anchor = deadline_row.start_dt or deadline_row.deadline_dt
    if deadline_row.frequency == "weekly":
        increment = timedelta(weeks=1)
    elif deadline_row.frequency == "biweekly":
        increment = timedelta(weeks=2)
    else:
        # "once" — the anchor is the only occurrence
        return anchor
    result = anchor
    now = datetime.utcnow()
    while result < now:
        result += increment
    return result


def web_search(query: str, max_results: int = 5) -> str:
    """
    Search the web using Tavily API and return formatted results.
    
    Args:
        query: Search query string
        max_results: Maximum number of results to return (default: 5)
    
    Returns:
        Formatted string with search results (title, content, URL) or error message
    """
    try:
        tavily_api_key = os.getenv("TAVILY_API_KEY")
        if not tavily_api_key:
            return "Error: TAVILY_API_KEY environment variable not set."
        
        client = TavilyClient(api_key=tavily_api_key)
        response = client.search(query=query, max_results=max_results)
        
        if not response.get("results"):
            return "No search results found."
        
        formatted_results = []
        for result in response["results"]:
            title = result.get("title", "No title")
            content = result.get("content", "No content available")
            url = result.get("url", "No URL available")
            formatted_results.append(f"Title: {title}\nContent: {content}\nSource: {url}\n")
        
        return "\n---\n".join(formatted_results)
    except Exception as e:
        return f"Error performing web search: {str(e)}"


def generate_ai_reply(group_id: str, question: str, db_group_id: int, username: str = None):
    db = SessionLocal()
    try:
        db_group = db.query(models.Group).filter(models.Group.id == db_group_id).first()
        if not db_group:
            return

        if not openai_client:
            error_message = "Error: OPENAI_API_KEY environment variable not set. Please configure your API key."
            if username:
                error_message = f"@{username}: {error_message}"
            ai_message = models.Message(
                group_id=db_group_id,
                user_id=None,
                content=error_message,
                is_AI=True
            )
            db.add(ai_message)
            db.commit()
            return

        try:
            # Fetch last 10 messages for conversation context
            recent_messages = db.query(models.Message).filter(
                models.Message.group_id == db_group_id
            ).order_by(models.Message.timestamp.desc()).limit(10).all()

            # Reverse to get chronological order (oldest first)
            recent_messages.reverse()

            # Build conversation history
            conversation_history = []
            for msg in recent_messages:
                if msg.is_AI:
                    conversation_history.append({
                        "role": "assistant",
                        "content": msg.content
                    })
                else:
                    msg_sender = msg.user.username if msg.user else "Unknown"
                    conversation_history.append({
                        "role": "user",
                        "content": f"{msg_sender}: {msg.content}"
                    })

            # The exact phrase Case 2 tells GPT-4o to say when the answer isn't
            # in the document — pipeline looks for this string to trigger Case 3
            REFUSAL_PHRASE = "I don't have that information in the uploaded documents"

            sources = []
            ai_text = ""

            # --- Filename detection: if the query names a specific uploaded file,
            # pull all its chunks directly and skip the cosine search entirely ---
            query_lower = question.lower()
            indexed_filenames = list_indexed_filenames(group_id)

            # Try matching on the human-readable original name first (e.g. "report.pdf"),
            # then fall back to safe_filename for docs indexed before original_filename was tracked
            matched_entry = next(
                (e for e in indexed_filenames if e["original_filename"].lower() in query_lower),
                None
            )
            if not matched_entry:
                matched_entry = next(
                    (e for e in indexed_filenames if e["safe_filename"].lower() in query_lower),
                    None
                )

            if matched_entry:
                matched_safe = matched_entry["safe_filename"]
                matched_display = matched_entry["original_filename"]
                print(f"[RAG] Filename match — serving all chunks for: {matched_safe}")
                file_chunks = get_chunks_by_filename(group_id, matched_safe)
                if file_chunks:
                    context_str = "\n\n---\n\n".join(file_chunks)
                    system_message = (
                        f"Answer the question using ONLY the document provided below. "
                        f"Cite the source inline as [{matched_display}] when you use information from it. "
                        f"Do not use general knowledge. Do not make up information.\n\n"
                        f"Document ({matched_display}):\n\n{context_str}"
                    )
                    messages = [{"role": "system", "content": system_message}]
                    messages.extend(conversation_history)
                    messages.append({"role": "user", "content": question})
                    response = openai_client.chat.completions.create(
                        model="gpt-4o",
                        messages=messages,
                        max_tokens=500,
                        temperature=0.7
                    )
                    ai_text = response.choices[0].message.content
                    sources = [{"type": "doc", "filename": matched_display}]

            if not ai_text:
                # --- Ask the RAG pipeline for the most relevant chunks and its confidence ---
                chunks, top_score = get_relevant_context(group_id, question)

                if top_score >= 0.0:
                    # --- Case 1: RAG path ---
                    # Top chunks scored well — answer directly from them
                    print(f"[RAG] Case 1 triggered — answering from document chunks")
                    context_parts = [
                        f"[{chunk['filename']}]\n{chunk['text']}"
                        for chunk in chunks
                    ]
                    context_str = "\n\n".join(context_parts)

                    system_message = (
                        "Answer the question using ONLY the context passages provided below. "
                        "Each passage is labelled with its source file in square brackets. "
                        "Cite the source inline as [filename.pdf] when you use information from it. "
                        "Do not use general knowledge. Do not make up information.\n\n"
                        "Context:\n\n" + context_str
                    )

                    messages = [{"role": "system", "content": system_message}]
                    messages.extend(conversation_history)
                    messages.append({"role": "user", "content": question})

                    response = openai_client.chat.completions.create(
                        model="gpt-4o",
                        messages=messages,
                        max_tokens=500,
                        temperature=0.7
                    )
                    ai_text = response.choices[0].message.content

                    sources = [
                        {"type": "doc", "filename": c["filename"], "chunk_index": c["chunk_index"]}
                        for c in chunks
                    ]

                else:
                    # --- Case 2: Full document fallback ---
                    # Chunks scored too low — rerank the best document's chunks
                    # against the query and pass the top hits instead of the full PDF
                    fallback_filename = get_top_document(chunks) if chunks else None
                    print(f"[RAG] Case 2 triggered — full document fallback: {fallback_filename}")
                    fallback_text = ""

                    if fallback_filename:
                        fallback_chunks = get_top_chunks_for_document(
                            group_id, question, fallback_filename
                        )
                        if fallback_chunks:
                            fallback_text = "\n\n---\n\n".join(fallback_chunks)

                    if fallback_text.strip():
                        system_message = (
                            f"Answer the question using ONLY the document passages provided below. "
                            f"Cite the source inline as [{fallback_filename}] when you use information from it. "
                            f"If the answer is genuinely not in the passages, reply with exactly: "
                            f"\"{REFUSAL_PHRASE}\"\n\n"
                            f"Passages from {fallback_filename}:\n\n{fallback_text}"
                        )

                        messages = [{"role": "system", "content": system_message}]
                        messages.extend(conversation_history)
                        messages.append({"role": "user", "content": question})

                        response = openai_client.chat.completions.create(
                            model="gpt-4o",
                            messages=messages,
                            max_tokens=500,
                            temperature=0.7
                        )
                        ai_text = response.choices[0].message.content

                        if REFUSAL_PHRASE not in ai_text:
                            # GPT-4o found the answer in the full document
                            sources = [{"type": "doc", "filename": fallback_filename}]
                        else:
                            # GPT-4o struck out — clear ai_text so we fall into Case 3
                            ai_text = ""
                            sources = []

                    if not ai_text:
                        # --- Case 3: Tavily web search ---
                        # Pipeline calls Tavily explicitly — results passed as plain
                        # text context, not as a tool to GPT-4o
                        print(f"[RAG] Case 3 triggered — Tavily web search fallback")
                        tavily_api_key = os.getenv("TAVILY_API_KEY")
                        web_results = []

                        if tavily_api_key:
                            try:
                                tavily_client_obj = TavilyClient(api_key=tavily_api_key)
                                tavily_response = tavily_client_obj.search(
                                    query=question, max_results=5
                                )
                                web_results = tavily_response.get("results", [])
                            except Exception:
                                web_results = []

                        if not web_results:
                            print(f"[RAG] All cases failed — returning final refusal")
                            ai_text = "I don't have enough information to answer that question."
                            sources = []
                        else:
                            context_parts = []
                            for result in web_results:
                                title = result.get("title", "")
                                content = result.get("content", "")
                                url = result.get("url", "")
                                context_parts.append(f"Title: {title}\n{content}\nSource: {url}")
                            context_str = "\n\n---\n\n".join(context_parts)

                            system_message = (
                                "The uploaded documents do not contain the answer to this question. "
                                "Answer using the web search results provided below. "
                                "Clearly state in your response that this information comes from "
                                "a web search, not the uploaded documents. "
                                "Include the source URLs in your response.\n\n"
                                "Web search results:\n\n" + context_str
                            )

                            messages = [{"role": "system", "content": system_message}]
                            messages.extend(conversation_history)
                            messages.append({"role": "user", "content": question})

                            response = openai_client.chat.completions.create(
                                model="gpt-4o",
                                messages=messages,
                                max_tokens=500,
                                temperature=0.7
                            )
                            ai_text = response.choices[0].message.content

                            sources = [
                                {"type": "web", "url": r.get("url", "")}
                                for r in web_results
                                if r.get("url")
                            ]

            if username:
                ai_text = f"@{username}: {ai_text}"

            ai_message = models.Message(
                group_id=db_group_id,
                user_id=None,
                content=ai_text,
                is_AI=True,
                sources=sources if sources else None
            )
            db.add(ai_message)
            db.commit()

        except Exception as e:
            # hide internal error details from the chat
            logging.error("AI reply failed for group %s: %s", db_group_id, str(e))
            error_message = "Error: Failed to get AI response. Please try again later."
            if username:
                error_message = f"@{username}: {error_message}"
            ai_message = models.Message(
                group_id=db_group_id,
                user_id=None,
                content=error_message,
                is_AI=True
            )
            db.add(ai_message)
            db.commit()

    except Exception as e:
        db.rollback()
        logging.error("Error saving AI reply: %s", str(e))
    finally:
        db.close()


@app.websocket("/ws/groups/{group_id}")
async def websocket_endpoint(websocket: WebSocket, group_id: str):
    # Token comes in as a query param since WS connections can't use headers
    token = websocket.query_params.get("token")
    if not token:
        await websocket.close(code=4008)
        return

    payload = decode_token(token)
    if not payload:
        await websocket.close(code=4008)
        return

    username = payload.get("sub") or payload.get("username")
    if not username:
        await websocket.close(code=4008)
        return

    db = SessionLocal()
    try:
        current_user = db.query(models.User).filter(models.User.username == username).first()
        if not current_user:
            await websocket.close(code=4008)
            return

        # Verify the user is actually a member of this group before accepting.
        if not check_group_access(group_id, current_user, db):
            await websocket.close(code=4008)
            return

        await websocket.accept()
        manager.connect(websocket, group_id, current_user.id)

        # Push any unread notifications that arrived while the user was offline
        pending = db.query(models.Notification, models.Group).join(
            models.Group, models.Group.id == models.Notification.group_id
        ).filter(
            models.Notification.recipient_id == current_user.id,
            models.Notification.is_read == False
        ).all()
        for notif, grp in pending:
            await manager.send_to_user(current_user.id, {
                "type": "notification",
                "id": notif.id,
                "sender_id": notif.sender_id,
                "message_id": notif.message_id,
                "group_id": notif.group_id,
                "group_string_id": grp.string_id,
                "created_at": notif.created_at.isoformat() + "Z"
            })

        # Look up the integer group PK once — needed for DB writes
        db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()

        try:
            while True:
                data = await websocket.receive_json()
                content = data.get("content", "").strip()
                if not content:
                    continue

                # Save the message
                new_message = models.Message(
                    group_id=db_group.id,
                    user_id=current_user.id,
                    content=content,
                    is_AI=False
                )
                db.add(new_message)
                db.commit()
                db.refresh(new_message)

                # Broadcast the message to everyone in the group
                await manager.broadcast_to_group(group_id, {
                    "type": "message",
                    "id": new_message.id,
                    "sender": current_user.username,
                    "text": content,
                    "is_bot": False,
                    "group_string_id": group_id,
                    "timestamp": new_message.timestamp.isoformat() + "Z"
                })

                # Nudge members who are connected to a *different* group's WS
                # so their badge counter updates without a full page refresh
                group_members = db.query(models.GroupMember).filter(
                    models.GroupMember.group_id == db_group.id
                ).all()
                watching = manager.active_user_ids_in_group(group_id)
                for member in group_members:
                    # sender already got the broadcast above
                    if member.user_id == current_user.id:
                        continue
                    # already got the broadcast because they're watching this group
                    if member.user_id in watching:
                        continue
                    await manager.send_to_user(member.user_id, {
                        "type": "message",
                        "id": new_message.id,
                        "sender": current_user.username,
                        "text": content,
                        "is_bot": False,
                        "group_string_id": group_id
                    })
                    # this user has no live WebSocket — send a push notification
                    # so they see a banner even if the app is backgrounded/closed
                    offline_tokens = [
                        pt.token for pt in db.query(models.PushToken).filter(
                            models.PushToken.user_id == member.user_id
                        ).all()
                    ]
                    if offline_tokens:
                        # run in a thread — requests.post is blocking and would
                        # stall the event loop (and all other WS connections) if awaited directly
                        asyncio.get_running_loop().run_in_executor(
                            None,
                            send_push_notifications,
                            offline_tokens,
                            f"{db_group.name}",
                            f"{current_user.username}: {content[:100]}",
                            {"groupId": group_id, "groupName": db_group.name},
                        )

                # Go through @mentions and notify the tagged users
                mentions = re.findall(r"@(\w+)", content)
                for mentioned_username in mentions:
                    # skip @ai — that's handled separately below
                    if mentioned_username.lower() == "ai":
                        continue
                    # skip self-mentions
                    if mentioned_username == current_user.username:
                        continue
                    # only notify if they're actually in this group
                    mentioned_user = db.query(models.User).join(
                        models.GroupMember,
                        models.GroupMember.user_id == models.User.id
                    ).filter(
                        models.User.username == mentioned_username,
                        models.GroupMember.group_id == db_group.id
                    ).first()
                    if not mentioned_user:
                        continue

                    notif = models.Notification(
                        recipient_id=mentioned_user.id,
                        sender_id=current_user.id,
                        message_id=new_message.id,
                        group_id=db_group.id
                    )
                    db.add(notif)
                    db.commit()
                    db.refresh(notif)

                    await manager.send_to_user(mentioned_user.id, {
                        "type": "notification",
                        "id": notif.id,
                        "sender_id": current_user.id,
                        "message_id": new_message.id,
                        "group_id": db_group.id,
                        "group_string_id": db_group.string_id,
                        "created_at": notif.created_at.isoformat() + "Z"
                    })

                # If @ai is in the message, run the RAG pipeline and broadcast the reply
                if "@ai" in content.lower():
                    question = re.sub(r"@ai\s*", "", content, flags=re.IGNORECASE).strip()
                    if question:
                        # generate_ai_reply is blocking (OpenAI call) — run it in a
                        # thread so the event loop stays free for other connections
                        loop = asyncio.get_running_loop()
                        await loop.run_in_executor(
                            None,
                            generate_ai_reply,
                            group_id,
                            question,
                            db_group.id,
                            current_user.username
                        )
                        # Fetch the AI message that was just written and broadcast it
                        ai_message = db.query(models.Message).filter(
                            models.Message.group_id == db_group.id,
                            models.Message.is_AI == True
                        ).order_by(models.Message.timestamp.desc()).first()
                        if ai_message:
                            await manager.broadcast_to_group(group_id, {
                                "type": "message",
                                "id": ai_message.id,
                                "sender": "AI Bot",
                                "text": ai_message.content,
                                "is_bot": True,
                                "group_string_id": group_id,
                                "timestamp": ai_message.timestamp.isoformat() + "Z"
                            })

        except WebSocketDisconnect:
            manager.disconnect(group_id, current_user.id, websocket)

    finally:
        db.close()


@app.websocket("/ws/home")
async def home_websocket_endpoint(websocket: WebSocket):
    # same auth pattern as /ws/groups/{group_id} — token comes in as a query param
    token = websocket.query_params.get("token")
    if not token:
        await websocket.close(code=4008)
        return

    payload = decode_token(token)
    if not payload:
        await websocket.close(code=4008)
        return

    username = payload.get("sub") or payload.get("username")
    if not username:
        await websocket.close(code=4008)
        return

    db = SessionLocal()
    try:
        current_user = db.query(models.User).filter(models.User.username == username).first()
        if not current_user:
            await websocket.close(code=4008)
            return

        await websocket.accept()
        # register under "__home__" so send_to_user can deliver cross-group nudges here
        manager.connect(websocket, "__home__", current_user.id)

        # send broadcasts from the last 7 days so the client has recent history on connect
        cutoff = datetime.utcnow() - timedelta(days=7)
        recent_broadcasts = (
            db.query(models.Broadcast)
            .filter(models.Broadcast.created_at >= cutoff)
            .order_by(models.Broadcast.created_at.asc())
            .all()
        )
        for b in recent_broadcasts:
            await websocket.send_text(json.dumps({
                "type": "broadcast",
                "id": b.id,
                "content": b.content,
                "sent_by": b.sent_by.username if b.sent_by else None,
                "created_at": b.created_at.isoformat() + "Z",
            }))

        try:
            while True:
                # client never sends on this socket — discard anything that arrives
                await websocket.receive_text()
        except WebSocketDisconnect:
            manager.disconnect("__home__", current_user.id, websocket)

    finally:
        db.close()


# Authentication dependency
def get_current_user(
    authorization: str = Header(None),
    db: Session = Depends(get_db)
) -> models.User:
    """
    Dependency to get the current authenticated user from JWT token.
    Expects Authorization header in format: "Bearer <token>"
    """
    if not authorization:
        raise HTTPException(status_code=401, detail="Authorization header missing")
    
    try:
        # Extract token from "Bearer <token>"
        scheme, token = authorization.split()
        if scheme.lower() != "bearer":
            raise HTTPException(status_code=401, detail="Invalid authentication scheme")
    except ValueError:
        raise HTTPException(status_code=401, detail="Invalid authorization header format")
    
    # Decode token
    payload = decode_token(token)
    if not payload:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    
    # Get username from token (assuming it's stored as "sub" or "username")
    username = payload.get("sub") or payload.get("username")
    if not username:
        raise HTTPException(status_code=401, detail="Token missing username")
    
    # Get user from database
    user = db.query(models.User).filter(models.User.username == username).first()
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    
    return user


@app.post("/auth/login")
def login(login_data: LoginRequest, db: Session = Depends(get_db)):
    """
    Login endpoint that verifies username and password, then returns a JWT token.
    """
    # Find user by username
    user = db.query(models.User).filter(models.User.username == login_data.username).first()
    
    if not user:
        raise HTTPException(status_code=401, detail="Invalid username or password")
    
    # Verify password
    if not verify_password(login_data.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid username or password")

    if not user.is_active:
        raise HTTPException(status_code=403, detail="Account is deactivated.")

    # Create access token with username in the payload
    access_token = create_access_token(data={"sub": user.username}, expires_minutes=60)
    
    return {
        "access_token": access_token,
        "token_type": "bearer"
    }


@app.post("/auth/microsoft")
def microsoft_login(body: MicrosoftLoginRequest, db: Session = Depends(get_db)):
    try:
        payload = verify_microsoft_token(body.id_token)
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid Microsoft token.")

    # Microsoft tokens put the email in "email" or fall back to "preferred_username"
    email = (payload.get("email") or payload.get("preferred_username", "")).lower().strip()

    if not email.lower().endswith("@e.ntu.edu.sg"):
        raise HTTPException(status_code=403, detail="Only NTU accounts are allowed.")

    user = db.query(models.User).filter(models.User.email == email).first()

    if not user:
        username = email.split("@")[0]
        user = models.User(
            username=username,
            email=email,
            role="student",
            # random UUID means the account has no usable password — must log in via Microsoft
            password_hash=hash_password(str(uuid.uuid4())),
            full_name=payload.get("name"),
            is_active=True,
        )
        db.add(user)
        db.commit()
        db.refresh(user)

    if not user.is_active:
        raise HTTPException(status_code=403, detail="Account is deactivated.")

    access_token = create_access_token(data={"sub": user.username})
    return {"access_token": access_token, "token_type": "bearer"}


@app.get("/auth/me")
def get_current_user_info(current_user: models.User = Depends(get_current_user)):
    """
    Get the current logged-in user's username and role.
    """
    return {
        "username": current_user.username,
        "role": current_user.role,
        "email": current_user.email,
        "full_name": current_user.full_name,
    }


@app.post("/auth/change-password")
def change_password(
    body: ChangePasswordRequest,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    raise HTTPException(status_code=403, detail="Password changes are not permitted. Contact an admin to reset your password.")
    if not verify_password(body.current_password, current_user.password_hash):
        raise HTTPException(status_code=403, detail="Current password is incorrect.")
    if len(body.new_password) < 8:
        raise HTTPException(status_code=400, detail="New password must be at least 8 characters.")
    current_user.password_hash = hash_password(body.new_password)
    db.commit()
    return {"message": "Password updated."}


class PushTokenRequest(BaseModel):
    token: str


@app.post("/push-token")
def register_push_token(
    body: PushTokenRequest,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Register or refresh an Expo push token for the current user."""
    existing = db.query(models.PushToken).filter(
        models.PushToken.token == body.token
    ).first()
    if not existing:
        db.add(models.PushToken(user_id=current_user.id, token=body.token))
        db.commit()
    elif existing.user_id != current_user.id:
        # token transferred to a new user (device re-login) — reassign it
        existing.user_id = current_user.id
        db.commit()
    return {"ok": True}


@app.get("/my-groups")
def get_my_groups(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get groups the current user is allowed to see.
    Returns groups with string IDs like "group-a" from database, sorted by group number.
    - Coordinator: sees all groups
    - Supervisor: sees only groups they are a member of
    - Student: sees only groups they are a member of
    """
    def extract_group_number(group_name):
        """Extract numeric part from group name (e.g., 'Group 1' -> 1)"""
        match = re.search(r'\d+', group_name)
        return int(match.group()) if match else 999
    
    if current_user.role in ("coordinator", "admin"):
        # Coordinator sees all groups
        all_groups = db.query(models.Group).all()
        # Sort by group number extracted from name
        sorted_groups = sorted(all_groups, key=lambda g: extract_group_number(g.name))
        return [{"id": group.string_id, "name": group.name} for group in sorted_groups]
    
    elif current_user.role == "supervisor" or current_user.role == "student":
        # Supervisor/Student sees groups they are a member of
        group_memberships = db.query(models.GroupMember).filter(
            models.GroupMember.user_id == current_user.id
        ).all()
        user_groups = [membership.group for membership in group_memberships]
        # Sort by group number extracted from name
        sorted_groups = sorted(user_groups, key=lambda g: extract_group_number(g.name))
        return [{"id": group.string_id, "name": group.name} for group in sorted_groups]
    
    else:
        raise HTTPException(status_code=403, detail="Unknown user role")


def check_group_access(group_id: str, current_user: models.User, db: Session) -> bool:
    """
    Check if the current user has access to the specified group.
    Returns True if access is allowed, False otherwise.
    """
    # Find group by string_id
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        return False

    # Coordinator and admin have access to all groups
    if current_user.role in ("coordinator", "admin"):
        return True
    
    # Check if user is a member of this group
    membership = db.query(models.GroupMember).filter(
        models.GroupMember.user_id == current_user.id,
        models.GroupMember.group_id == db_group.id
    ).first()
    
    return membership is not None


def check_summary_access(group_id: str, current_user: models.User, db: Session) -> bool:
    """
    Check if the current user has access to summaries for the specified group.
    Returns True if access is allowed, False otherwise.
    Access rules: coordinator OR (member of group - includes supervisors and students).
    """
    # Find group by string_id
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        return False
    
    # Coordinator and admin have access to all groups
    if current_user.role in ("coordinator", "admin"):
        return True

    # All group members (supervisors and students) have access
    membership = db.query(models.GroupMember).filter(
        models.GroupMember.user_id == current_user.id,
        models.GroupMember.group_id == db_group.id
    ).first()
    return membership is not None


@app.get("/groups/{group_id}/members")
def list_group_members(
    group_id: str,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")

    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")

    # everyone in the group except the caller — they can't @mention themselves
    members = db.query(models.User).join(
        models.GroupMember,
        models.GroupMember.user_id == models.User.id
    ).filter(
        models.GroupMember.group_id == db_group.id,
        models.User.id != current_user.id
    ).all()

    result = [{"username": m.username, "full_name": m.full_name} for m in members]
    # "ai" is a special mention target that triggers the RAG pipeline
    result.append({"username": "ai"})

    return result


@app.get("/groups/{group_id}/messages")
def list_messages(
    group_id: str,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # Find group by string_id
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")
    
    # Check authorization
    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")
    
    # Query messages from database
    messages = db.query(models.Message).filter(
        models.Message.group_id == db_group.id
    ).order_by(models.Message.timestamp).all()
    
    # Convert to API format: [{id, sender, text, is_bot, message_type, image_url, timestamp}]
    result = []
    for msg in messages:
        sender = "AI Bot" if msg.is_AI else (msg.user.username if msg.user else "Unknown")
        msg_type = msg.message_type if msg.message_type else "text"
        image_url = (
            f"/groups/{group_id}/messages/{msg.id}/image"
            if msg_type == "image" else None
        )
        result.append({
            "id": msg.id,
            "sender": sender,
            "text": msg.content,
            "is_bot": msg.is_AI,
            "message_type": msg_type,
            "image_url": image_url,
            "timestamp": msg.timestamp.isoformat() + "Z"
        })

    return result


@app.post("/groups/{group_id}/messages")
def add_message(
    group_id: str,
    message: NewMessage,
    background_tasks: BackgroundTasks,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # Find group by string_id
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")
    
    # Check authorization
    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")

    # Save the user message to database
    new_message = models.Message(
        group_id=db_group.id,
        user_id=current_user.id,
        content=message.text,
        is_AI=False
    )
    db.add(new_message)
    db.commit()
    db.refresh(new_message)

    # Check if message starts with "@ai" and schedule AI response in background
    if message.text.startswith("@ai"):
        # Extract the question (remove "@ai" prefix and strip whitespace)
        question = message.text[3:].strip()
        
        if question:
            # Schedule AI reply generation as a background task
            # Pass the username so the AI can mention them in the response
            background_tasks.add_task(generate_ai_reply, group_id, question, db_group.id, current_user.username)

    # Return in API format: {id, sender, text, is_bot}
    return {
        "id": new_message.id,
        "sender": current_user.username,
        "text": new_message.content,
        "is_bot": False
    }


@app.post("/groups/{group_id}/messages/image")
async def upload_image_message(
    group_id: str,
    file: UploadFile = File(...),
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")

    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")

    # only JPEG and PNG accepted
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in (".jpg", ".jpeg", ".png"):
        raise HTTPException(status_code=400, detail="Only JPEG and PNG images are allowed")

    content = await file.read()

    # 5 MB hard limit
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Image must be 5 MB or smaller")

    # JPEG and PNG magic bytes
    is_jpeg = content[:3] == b"\xff\xd8\xff"
    is_png = content[:8] == b"\x89PNG\r\n\x1a\n"
    if not (is_jpeg or is_png):
        raise HTTPException(status_code=400, detail="File content does not match a valid JPEG or PNG image")

    # strip any directory components from the client-supplied filename
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = f"{ts}_{os.path.basename(file.filename)}"

    if is_blob_storage_enabled():
        blob_name = f"{group_id}/{safe_filename}"
        try:
            upload_blob("images", blob_name, content)
        except Exception as e:
            logging.error("Failed to upload image to blob storage: %s", str(e))
            raise HTTPException(status_code=500, detail="Failed to save image.")
        stored_content = blob_name
    else:
        file_path = IMAGE_STORAGE_DIR / group_id / safe_filename
        file_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            with open(file_path, "wb") as f:
                f.write(content)
        except Exception as e:
            logging.error("Failed to save image: %s", str(e))
            raise HTTPException(status_code=500, detail="Failed to save image.")
        stored_content = str(file_path)

    new_message = models.Message(
        group_id=db_group.id,
        user_id=current_user.id,
        content=stored_content,
        is_AI=False,
        message_type="image"
    )
    db.add(new_message)
    db.commit()
    db.refresh(new_message)

    image_url = f"/groups/{group_id}/messages/{new_message.id}/image"

    await manager.broadcast_to_group(group_id, {
        "type": "message",
        "id": new_message.id,
        "sender": current_user.username,
        "text": "",
        "is_bot": False,
        "message_type": "image",
        "image_url": image_url,
        "group_string_id": group_id,
        "timestamp": new_message.timestamp.isoformat()
    })

    return {
        "id": new_message.id,
        "sender": current_user.username,
        "message_type": "image",
        "image_url": image_url,
        "timestamp": new_message.timestamp.isoformat()
    }


@app.get("/groups/{group_id}/messages/{message_id}/image")
def get_image_message(
    group_id: str,
    message_id: int,
    token: str = Query(None),        # Image component can't set headers — accept token here too
    authorization: str = Header(None),
    db: Session = Depends(get_db)
):
    # prefer query-param token (used by the React Native Image component), fall back to header
    raw_token = token
    if not raw_token and authorization:
        try:
            _, raw_token = authorization.split()
        except ValueError:
            raw_token = None

    if not raw_token:
        raise HTTPException(status_code=401, detail="Authorization required")

    payload = decode_token(raw_token)
    if not payload:
        raise HTTPException(status_code=401, detail="Invalid or expired token")

    username = payload.get("sub") or payload.get("username")
    current_user = db.query(models.User).filter(models.User.username == username).first()
    if not current_user:
        raise HTTPException(status_code=401, detail="User not found")

    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")

    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")

    msg = db.query(models.Message).filter(
        models.Message.id == message_id,
        models.Message.group_id == db_group.id,
        models.Message.message_type == "image"
    ).first()
    if not msg:
        raise HTTPException(status_code=404, detail="Image message not found")

    if is_blob_storage_enabled():
        try:
            data = download_blob("images", msg.content)
        except FileNotFoundError:
            raise HTTPException(status_code=404, detail="Image file not found in blob storage")
        ext = os.path.splitext(msg.content)[1].lower()
        mime = "image/jpeg" if ext in (".jpg", ".jpeg") else "image/png"
        return Response(content=data, media_type=mime)
    else:
        # resolve() collapses any .. or symlinks so is_relative_to() can't be tricked
        file_path = Path(msg.content).resolve()
        if not file_path.is_relative_to(IMAGE_STORAGE_DIR.resolve()):
            raise HTTPException(status_code=404, detail="Image file not found on disk")
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Image file not found on disk")

        ext = file_path.suffix.lower()
        mime = "image/jpeg" if ext in (".jpg", ".jpeg") else "image/png"
        return FileResponse(path=file_path, media_type=mime)


@app.get("/groups/{group_id}/documents")
def list_documents(
    group_id: str,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List all PDF documents for a group"""
    # Check if group exists in database
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")
    
    # Check authorization
    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")
    
    # Query documents from database
    documents = db.query(models.Document).filter(
        models.Document.group_id == group_id
    ).order_by(models.Document.created_at.desc()).all()
    
    # Convert to API format: [{id, filename, uploaded_at, file_path, uploaded_by, file_size}]
    import os
    result = []
    for doc in documents:
        # Get file size
        file_size = 0
        if os.path.exists(doc.stored_path):
            file_size = os.path.getsize(doc.stored_path)
        
        # Get uploader username
        uploader_name = "Unknown"
        if doc.uploaded_by:
            uploader_name = doc.uploaded_by.username
        
        # Ensure UTC timestamp is marked with 'Z' suffix
        uploaded_at_iso = doc.created_at.isoformat()
        if not uploaded_at_iso.endswith('Z'):
            uploaded_at_iso = uploaded_at_iso + 'Z'
        
        result.append({
            "id": doc.id,
            "filename": doc.filename,
            "uploaded_at": uploaded_at_iso,
            "uploaded_by": uploader_name,
            "file_size": file_size,
        })
    
    return result


@app.post("/groups/{group_id}/documents")
async def upload_document(
    group_id: str,
    file: UploadFile = File(...),
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload a PDF document for a group"""
    # Check if group exists in database
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")
    
    # Check authorization
    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")
    
    # Validate file type
    allowed_extensions = ['.pdf', '.doc', '.docx']
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Only PDF, DOC, and DOCX files are allowed")
    
    # Generate unique filename to avoid conflicts.
    # strip any directory components from the client-supplied filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = f"{timestamp}_{os.path.basename(file.filename)}"

    # Read content first so we can check size before touching disk
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File must be 20 MB or smaller")

    # PDF, DOCX (ZIP), and legacy DOC magic bytes
    is_pdf = content[:4] == b"%PDF"
    is_docx = content[:2] == b"PK"
    is_doc = content[:4] == b"\xd0\xcf\x11\xe0"
    if file_ext == ".pdf" and not is_pdf:
        raise HTTPException(status_code=400, detail="File content does not match a valid PDF")
    if file_ext == ".docx" and not is_docx:
        raise HTTPException(status_code=400, detail="File content does not match a valid DOCX file")
    if file_ext == ".doc" and not is_doc:
        raise HTTPException(status_code=400, detail="File content does not match a valid DOC file")

    if is_blob_storage_enabled():
        blob_name = f"{group_id}/{safe_filename}"
        try:
            upload_blob("documents", blob_name, content)
        except Exception as e:
            logging.error("Failed to upload to blob storage: %s", str(e))
            raise HTTPException(status_code=500, detail="Failed to save file.")
        stored_path = blob_name
    else:
        file_path = PDF_STORAGE_DIR / group_id / safe_filename
        file_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            with open(file_path, "wb") as f:
                f.write(content)
        except Exception as e:
            logging.error("Failed to save file: %s", str(e))
            raise HTTPException(status_code=500, detail="Failed to save file.")
        stored_path = str(file_path)

    # Store metadata in database
    document = models.Document(
        group_id=group_id,
        uploaded_by_user_id=current_user.id,
        filename=file.filename,
        stored_path=stored_path,
        created_at=datetime.utcnow()
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    # read from memory instead of disk so this works for both storage backends
    # .doc is the old binary Word format — python-docx can't read it, so those are stored but not indexed.
    if file_ext == '.pdf':
        try:
            reader = PdfReader(io.BytesIO(content))
            pdf_text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if pdf_text.strip():
                # Use safe_filename (timestamped) so two uploads of "report.pdf" don't
                # clobber each other's ChromaDB embeddings.
                index_document(group_id, safe_filename, pdf_text, original_filename=file.filename)
        except Exception as e:
            # Don't block the upload if indexing fails — just log it
            logging.warning("Failed to index %s for RAG: %s", file.filename, str(e))
    elif file_ext == '.docx':
        try:
            reader = DocxDocument(io.BytesIO(content))
            doc_text = "\n".join(p.text for p in reader.paragraphs if p.text.strip())
            if doc_text.strip():
                index_document(group_id, safe_filename, doc_text, original_filename=file.filename)
        except Exception as e:
            logging.warning("Failed to index %s for RAG: %s", file.filename, str(e))
    # .doc files land here — stored, skipped for indexing since python-docx can't parse the old binary format

    file_size = len(content)

    # Ensure UTC timestamp is marked with 'Z' suffix
    uploaded_at_iso = document.created_at.isoformat()
    if not uploaded_at_iso.endswith('Z'):
        uploaded_at_iso = uploaded_at_iso + 'Z'

    return {
        "id": document.id,
        "filename": document.filename,
        "uploaded_at": uploaded_at_iso,
        "uploaded_by": current_user.username,
        "file_size": file_size,
    }


@app.get("/groups/{group_id}/documents/{document_id}")
def download_document(
    group_id: str,
    document_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Download a PDF document"""
    # Check if group exists in database
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")
    
    # Check authorization
    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")
    
    # Find the document in database
    document = db.query(models.Document).filter(
        models.Document.id == document_id,
        models.Document.group_id == group_id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if is_blob_storage_enabled():
        try:
            data = download_blob("documents", document.stored_path)
        except FileNotFoundError:
            raise HTTPException(status_code=404, detail="File not found in blob storage")
        return Response(
            content=data,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{document.filename}"'}
        )
    else:
        file_path = Path(document.stored_path)
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="File not found on disk")
        return FileResponse(
            path=file_path,
            filename=document.filename,
            media_type="application/pdf"
        )


@app.delete("/groups/{group_id}/documents/{document_id}")
def delete_document(
    group_id: str,
    document_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a document from the database and file system"""
    # Check if group exists in database
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")
    
    # Check authorization
    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")
    
    # Find the document in database
    document = db.query(models.Document).filter(
        models.Document.id == document_id,
        models.Document.group_id == group_id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # Students can only delete their own uploads; supervisors and coordinators can delete any.
    is_uploader = document.uploaded_by_user_id == current_user.id
    is_privileged = current_user.role in ("supervisor", "coordinator", "admin")
    if not (is_uploader or is_privileged):
        raise HTTPException(status_code=403, detail="Only the uploader, a supervisor, coordinator, or admin can delete this document")

    if is_blob_storage_enabled():
        delete_blob("documents", document.stored_path)
    else:
        # Path-traversal guard — make sure stored_path is still inside our uploads dir
        file_path = Path(document.stored_path).resolve()
        if not file_path.is_relative_to(PDF_STORAGE_DIR.resolve()):
            logging.warning("Skipping deletion of file outside storage directory: %s", file_path)
        elif file_path.exists():
            try:
                file_path.unlink()
            except Exception as e:
                logging.error("Error deleting file %s: %s", file_path, str(e))

    # Delete the document from database
    db.delete(document)
    db.commit()

    return {"message": "Document deleted successfully"}


@app.get("/groups/{group_id}/summary")
def get_summary(
    group_id: str,
    range: str = Query("weekly", description="Summary range type (e.g., 'weekly', 'full')"),
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get the latest summary for a group.
    Access: coordinator OR (member of group - includes supervisors and students).
    """
    # Check if group exists in database
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")
    
    # Check authorization (coordinator or supervisor only, no students)
    if not check_summary_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")

    if range not in ("weekly", "full"):
        raise HTTPException(status_code=400, detail="range must be 'weekly' or 'full'")

    # Query for the latest summary matching the group_id and range_type
    summary = db.query(models.Summary).filter(
        models.Summary.group_id == group_id,
        models.Summary.range_type == range
    ).order_by(models.Summary.created_at.desc()).first()
    
    # If no summary exists, return empty summary structure
    if not summary:
        return {
            "group_id": group_id,
            "range": range,
            "summary_text": "",
            "created_at": None,
            "start_time": None,
            "end_time": None,
            "source_last_message_ts": None,
            "source_message_count": None
        }
    
    # Return summary data
    return {
        "group_id": summary.group_id,
        "range": summary.range_type,
        "summary_text": summary.summary_text,
        "created_at": summary.created_at.isoformat() + "Z" if summary.created_at else None,
        "start_time": summary.start_time.isoformat() + "Z" if summary.start_time else None,
        "end_time": summary.end_time.isoformat() + "Z" if summary.end_time else None,
        "source_last_message_ts": summary.source_last_message_ts.isoformat() + "Z" if summary.source_last_message_ts else None,
        "source_message_count": summary.source_message_count
    }


@app.post("/groups/{group_id}/summary")
def generate_summary(
    group_id: str,
    range: str = Query("weekly", description="Summary range type (e.g., 'weekly', 'full')"),
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Generate and save a summary for a group.
    Access: coordinator OR (member of group - includes supervisors and students).
    """
    # Check if group exists in database
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")
    
    # Check authorization (coordinator or group member)
    if not check_summary_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")

    if range not in ("weekly", "full"):
        raise HTTPException(status_code=400, detail="range must be 'weekly' or 'full'")

    # Handle different range types
    if range == "weekly":
        # Get messages from the last 7 days
        cutoff_time = datetime.utcnow() - timedelta(days=7)
        messages = db.query(models.Message).filter(
            models.Message.group_id == db_group.id,
            models.Message.timestamp >= cutoff_time
        ).order_by(models.Message.timestamp.asc()).all()
    else:
        # For other range types, fetch all messages (can be extended later)
        messages = db.query(models.Message).filter(
            models.Message.group_id == db_group.id
        ).order_by(models.Message.timestamp.asc()).all()
    
    # Check if there are any messages
    if not messages:
        # Return friendly message
        return {
            "group_id": group_id,
            "range": range,
            "summary_text": "No messages in the selected period.",
            "created_at": datetime.utcnow().isoformat() + "Z",
            "start_time": None,
            "end_time": None,
            "source_last_message_ts": None,
            "source_message_count": 0
        }
    
    # Get the latest message timestamp for optimization check
    latest_message_ts = messages[-1].timestamp
    
    # Check if we can reuse an existing summary (optimization)
    existing_summary = db.query(models.Summary).filter(
        models.Summary.group_id == group_id,
        models.Summary.range_type == range
    ).order_by(models.Summary.created_at.desc()).first()

    if existing_summary and existing_summary.source_last_message_ts:
        if existing_summary.source_last_message_ts == latest_message_ts:
            return {
                "group_id": existing_summary.group_id,
                "range": existing_summary.range_type,
                "summary_text": existing_summary.summary_text,
                "created_at": existing_summary.created_at.isoformat() + "Z" if existing_summary.created_at else None,
                "start_time": existing_summary.start_time.isoformat() + "Z" if existing_summary.start_time else None,
                "end_time": existing_summary.end_time.isoformat() + "Z" if existing_summary.end_time else None,
                "source_last_message_ts": existing_summary.source_last_message_ts.isoformat() + "Z" if existing_summary.source_last_message_ts else None,
                "source_message_count": existing_summary.source_message_count
            }

    # Build transcript from messages.
    # Image messages get a placeholder in the text transcript and are base64-encoded
    # separately so they can be passed to the vision model in step 2.
    transcript_lines = []
    image_blocks = []  # {mime, data} dicts in conversation order
    for msg in messages:
        timestamp_str = msg.timestamp.strftime("%Y-%m-%d %H:%M")
        username = "AI Bot" if msg.is_AI else (msg.user.username if msg.user else "Unknown")
        msg_type = msg.message_type if msg.message_type else "text"

        if msg_type == "image":
            # mark where the image appeared in the conversation
            transcript_lines.append(f"[{timestamp_str}] {username}: [sent an image]")

            if is_blob_storage_enabled():
                try:
                    img_bytes = download_blob("images", msg.content)
                    ext = os.path.splitext(msg.content)[1].lower()
                    mime = "image/jpeg" if ext in (".jpg", ".jpeg") else "image/png"
                    b64 = base64.b64encode(img_bytes).decode("utf-8")
                    image_blocks.append({"mime": mime, "data": b64})
                except Exception as e:
                    logging.warning("Could not read image %s: %s", msg.content, str(e))
            else:
                # only read files that live inside this group's image folder
                # resolve() collapses any .. so a crafted DB value can't escape the group dir
                file_path = Path(msg.content).resolve()
                if not file_path.is_relative_to((IMAGE_STORAGE_DIR / group_id).resolve()):
                    logging.warning("Skipping image outside group directory: %s", msg.content)
                    continue

                if not file_path.exists():
                    logging.warning("Skipping missing image file: %s", msg.content)
                    continue

                try:
                    ext = file_path.suffix.lower()
                    mime = "image/jpeg" if ext in (".jpg", ".jpeg") else "image/png"
                    with open(file_path, "rb") as f:
                        b64 = base64.b64encode(f.read()).decode("utf-8")
                    image_blocks.append({"mime": mime, "data": b64})
                except Exception as e:
                    logging.warning("Could not read image %s: %s", msg.content, str(e))
        else:
            transcript_lines.append(f"[{timestamp_str}] {username}: {msg.content}")

    transcript = "\n".join(transcript_lines)

    # use the last AI summary as context so the new one can note what's changed
    previous_ai_text = existing_summary.summary_text if existing_summary else None

    # pull the most recent student-written summary for comparison context
    latest_student = db.query(models.StudentSummary).filter(
        models.StudentSummary.group_id == group_id
    ).order_by(models.StudentSummary.created_at.desc()).first()
    student_summary_text = latest_student.summary_text if latest_student else None

    # Check if OpenAI client is available
    if not openai_client:
        # If OpenAI is not available, save an error summary
        summary_text = "Error: OPENAI_API_KEY environment variable not set. Please configure your API key."
    else:
        try:
            # Build prompt for simple summary generation
            system_prompt = (
                "You are a helpful assistant that creates concise, readable summaries of group chat conversations. "
                "Analyze the conversation transcript and provide a summary in EXACTLY this format:\n\n"
                "First, write a short plain-language paragraph (2-3 sentences) explaining what happened in the group recently. "
                "Then, on a new line, write 'Key points:' followed by 3-6 concise bullet points on separate lines. "
                "Finally, on a new line, write 'Supervisor Action Plan:' followed by 2-3 actionable recommendations for supervisors based on the conversation.\n\n"
                "CRITICAL FORMATTING RULES:\n"
                "- DO NOT use any markdown headings (no #, ##, ###, etc.)\n"
                "- DO NOT use markdown bold (**text**)\n"
                "- DO NOT create sections like 'Highlights', 'Decisions', 'Open Questions', or 'Action Items'\n"
                "- Use simple dashes (-) or asterisks (*) for bullets, NOT markdown\n"
                "- Write in plain text only\n"
                "- Keep the paragraph conversational and easy to read\n"
                "- Make bullets action-oriented and skimmable\n"
                "- Supervisor Action Plan bullets should be concrete, specific, and actionable\n"
                "- Example format:\n"
                "  The group discussed project timelines and resource allocation. Several team members shared updates on their progress.\n\n"
                "  Key points:\n"
                "  - Project deadline moved to next month\n"
                "  - Need to assign additional developer\n"
                "  - Client feedback session scheduled for Friday\n\n"
                "  Supervisor Action Plan:\n"
                "  - Schedule a meeting with the team to discuss resource allocation and clarify roles\n"
                "  - Review the project timeline and provide guidance on priority tasks\n"
                "  - Follow up with the client to confirm feedback session details\n\n"
                "For any images shared in the conversation, include a brief description of what the image shows as part of the summary. "
                "Do not say the image lacks context — describe what you can see."
            )

            # if a previous AI summary exists, ask the model to note what has changed
            if previous_ai_text:
                system_prompt += (
                    "\n\nFor reference, here is the previous AI summary generated for this group:\n\n"
                    + previous_ai_text
                    + "\n\nIn your new summary, briefly note what has changed or progressed since then."
                )

            # if the students wrote their own summary, ask the model to compare it
            if student_summary_text:
                system_prompt += (
                    "\n\nThe students have also written their own account of their progress:\n\n"
                    + student_summary_text
                    + "\n\nCompare what the students say with what the chat transcript shows. "
                    "Note any differences or areas of growth, and incorporate this into the Supervisor Action Plan."
                )

            # text first, then images in the order they appeared in the chat
            user_content = [
                {
                    "type": "text",
                    "text": f"Please create a summary of the following group chat conversation:\n\n{transcript}"
                }
            ]
            for img in image_blocks:
                user_content.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{img['mime']};base64,{img['data']}"
                    }
                })

            response = openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content}
                ],
                max_tokens=1200,
                temperature=0.7
            )
            
            summary_text = response.choices[0].message.content
        except Exception as e:
            # Handle API errors gracefully
            summary_text = f"Error: Failed to generate summary. {str(e)}"
    
    # Calculate time range
    start_time = messages[0].timestamp
    end_time = messages[-1].timestamp
    
    # Save summary to database
    new_summary = models.Summary(
        group_id=group_id,
        range_type=range,
        start_time=start_time,
        end_time=end_time,
        summary_text=summary_text,
        created_by_user_id=current_user.id,
        source_last_message_ts=latest_message_ts,
        source_message_count=len(messages),
        created_at=datetime.utcnow()
    )
    db.add(new_summary)
    db.commit()
    db.refresh(new_summary)
    
    # Return summary data
    return {
        "group_id": new_summary.group_id,
        "range": new_summary.range_type,
        "summary_text": new_summary.summary_text,
        "created_at": new_summary.created_at.isoformat() + "Z" if new_summary.created_at else None,
        "start_time": new_summary.start_time.isoformat() + "Z" if new_summary.start_time else None,
        "end_time": new_summary.end_time.isoformat() + "Z" if new_summary.end_time else None,
        "source_last_message_ts": new_summary.source_last_message_ts.isoformat() + "Z" if new_summary.source_last_message_ts else None,
        "source_message_count": new_summary.source_message_count
    }


@app.get("/groups/{group_id}/summary/history")
def get_summary_history(
    group_id: str,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Return all AI summaries for a group, newest first.
    Access: coordinator OR group member (same as the summary endpoints).
    """
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")

    if not check_summary_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")

    rows = db.query(models.Summary).filter(
        models.Summary.group_id == group_id
    ).order_by(models.Summary.created_at.desc()).all()

    return [
        {
            "id": r.id,
            "summary_text": r.summary_text,
            "created_at": r.created_at.isoformat() + "Z",
            "source_message_count": r.source_message_count,
        }
        for r in rows
    ]


@app.get("/groups/{group_id}/student-summary")
def get_student_summary(
    group_id: str,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get the most recent student summary for a group.
    Access: All group members (students, supervisors, coordinators).
    """
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")

    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")

    # most recent entry, or empty string if none exists yet
    row = db.query(models.StudentSummary).filter(
        models.StudentSummary.group_id == group_id
    ).order_by(models.StudentSummary.created_at.desc()).first()

    return {
        "group_id": group_id,
        "summary_text": row.summary_text if row else "",
        "ai_summary_copy": row.ai_summary_copy if row else None,
    }


@app.post("/groups/{group_id}/student-summary")
def update_student_summary(
    group_id: str,
    request: StudentSummaryRequest,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Save a new student summary for a group (inserts a row, keeps history).
    Access: All group members (students, supervisors, coordinators).
    """
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")

    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")

    # insert a new row so the full edit history is preserved
    new_entry = models.StudentSummary(
        group_id=group_id,
        summary_text=request.summary_text,
        ai_summary_copy=request.ai_summary_copy,
        created_by_user_id=current_user.id,
    )
    db.add(new_entry)
    db.commit()
    db.refresh(new_entry)

    return {
        "group_id": group_id,
        "summary_text": new_entry.summary_text,
        "ai_summary_copy": new_entry.ai_summary_copy,
    }


@app.get("/groups/{group_id}/student-summary/history")
def get_student_summary_history(
    group_id: str,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Return all saved student summaries for a group, newest first.
    Access: All group members (students, supervisors, coordinators).
    """
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")

    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")

    rows = db.query(models.StudentSummary).filter(
        models.StudentSummary.group_id == group_id
    ).order_by(models.StudentSummary.created_at.desc()).all()

    return [
        {
            "id": r.id,
            "summary_text": r.summary_text,
            "ai_summary_copy": r.ai_summary_copy,
            "is_submitted": r.is_submitted,
            "submitted_at": r.submitted_at.isoformat() + "Z" if r.submitted_at else None,
            "created_at": r.created_at.isoformat() + "Z",
        }
        for r in rows
    ]


@app.post("/groups/{group_id}/student-summary/submit")
def submit_student_summary(
    group_id: str,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role != "student":
        raise HTTPException(status_code=403, detail="Only students can submit a summary")

    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")

    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")

    # most recent save is the one that gets submitted
    latest = db.query(models.StudentSummary).filter(
        models.StudentSummary.group_id == group_id
    ).order_by(models.StudentSummary.created_at.desc()).first()

    if not latest:
        raise HTTPException(status_code=404, detail="No summary to submit")

    # Check the deadline and decide whether to block or flag as late
    deadline_row = db.query(models.Deadline).order_by(models.Deadline.created_at.desc()).first()
    is_late = False
    if deadline_row:
        next_dt = compute_next_deadline(deadline_row)
        if datetime.utcnow() > next_dt:
            if deadline_row.is_hard:
                raise HTTPException(status_code=403, detail="Submission deadline has passed.")
            else:
                is_late = True

    latest.is_submitted = True
    latest.submitted_at = datetime.utcnow()
    latest.is_late = is_late
    db.commit()
    db.refresh(latest)

    return {
        "id": latest.id,
        "group_id": group_id,
        "summary_text": latest.summary_text,
        "ai_summary_copy": latest.ai_summary_copy,
        "is_submitted": latest.is_submitted,
        "submitted_at": latest.submitted_at.isoformat() + "Z",
        "is_late": latest.is_late,
        "created_at": latest.created_at.isoformat() + "Z",
    }


@app.get("/deadline")
def get_deadline(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    row = db.query(models.Deadline).order_by(models.Deadline.created_at.desc()).first()
    if not row:
        return None
    next_dt = compute_next_deadline(row)
    return {
        "next_deadline_dt": next_dt.isoformat() + "Z",
        "frequency": row.frequency,
        "is_hard": row.is_hard,
        "set_by": row.set_by.username if row.set_by else None,
    }


@app.post("/deadline")
def set_deadline(
    body: DeadlineRequest,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role not in ("coordinator", "admin"):
        raise HTTPException(status_code=403, detail="Only coordinators can set the deadline")

    if body.frequency not in ("once", "weekly", "biweekly"):
        raise HTTPException(status_code=400, detail="frequency must be 'once', 'weekly', or 'biweekly'")

    new_deadline = models.Deadline(
        deadline_dt=body.start_dt,  # keep deadline_dt in sync for backward compat
        start_dt=body.start_dt,
        frequency=body.frequency,
        is_hard=body.is_hard,
        set_by_user_id=current_user.id,
    )
    db.add(new_deadline)
    db.commit()
    db.refresh(new_deadline)

    next_dt = compute_next_deadline(new_deadline)
    return {
        "next_deadline_dt": next_dt.isoformat() + "Z",
        "frequency": new_deadline.frequency,
        "is_hard": new_deadline.is_hard,
        "set_by": current_user.username,
    }


@app.get("/course-period")
def get_course_period(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    row = db.query(models.CoursePeriod).order_by(models.CoursePeriod.created_at.desc()).first()
    if not row:
        return None
    return {
        "start_date": row.start_date.isoformat(),
        "end_date": row.end_date.isoformat(),
        "set_by": row.set_by.username if row.set_by else None,
    }


@app.post("/course-period")
def set_course_period(
    body: CoursePeriodRequest,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role not in ("coordinator", "admin"):
        raise HTTPException(status_code=403, detail="Only coordinators can set the course period")

    if body.end_date <= body.start_date:
        raise HTTPException(status_code=400, detail="end_date must be after start_date")

    new_period = models.CoursePeriod(
        start_date=body.start_date,
        end_date=body.end_date,
        set_by_user_id=current_user.id,
    )
    db.add(new_period)
    db.commit()
    db.refresh(new_period)

    return {
        "start_date": new_period.start_date.isoformat(),
        "end_date": new_period.end_date.isoformat(),
        "set_by": current_user.username,
    }


@app.get("/notifications")
def get_notifications(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    results = db.query(models.Notification, models.Group).join(
        models.Group, models.Group.id == models.Notification.group_id
    ).filter(
        models.Notification.recipient_id == current_user.id,
        models.Notification.is_read == False
    ).all()
    return [
        {
            "id": n.id,
            "sender_id": n.sender_id,
            "message_id": n.message_id,
            "group_id": n.group_id,
            "group_string_id": g.string_id,
            "created_at": n.created_at.isoformat() + "Z"
        }
        for n, g in results
    ]


@app.post("/notifications/{notification_id}/read")
def mark_notification_read(
    notification_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    notif = db.query(models.Notification).filter(
        models.Notification.id == notification_id
    ).first()
    # 404 if missing or belongs to a different user
    if not notif or notif.recipient_id != current_user.id:
        raise HTTPException(status_code=404, detail="Notification not found")
    notif.is_read = True
    db.commit()
    return {"status": "ok"}


@app.post("/broadcast")
async def send_broadcast(
    body: BroadcastRequest,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role not in ("coordinator", "admin"):
        raise HTTPException(status_code=403, detail="Coordinator or admin access required.")
    if not body.content.strip():
        raise HTTPException(status_code=400, detail="Content cannot be empty.")
    if len(body.content) > 1000:
        raise HTTPException(status_code=400, detail="Content must be 1000 characters or fewer.")
    broadcast = models.Broadcast(content=body.content, sent_by_user_id=current_user.id)
    db.add(broadcast)
    db.commit()
    db.refresh(broadcast)
    payload = {
        "type": "broadcast",
        "id": broadcast.id,
        "content": broadcast.content,
        "sent_by": current_user.username,
        "created_at": broadcast.created_at.isoformat() + "Z",
    }
    all_user_ids: set[int] = set()
    for group_connections in list(manager.active_connections.values()):
        all_user_ids.update(group_connections.keys())
    for uid in all_user_ids:
        await manager.send_to_user(uid, payload)
    return payload


@app.get("/broadcasts")
def get_broadcasts(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    rows = (
        db.query(models.Broadcast)
        .order_by(models.Broadcast.created_at.desc())
        .limit(10)
        .all()
    )
    return [
        {
            "id": b.id,
            "content": b.content,
            "sent_by": b.sent_by.username if b.sent_by else None,
            "created_at": b.created_at.isoformat() + "Z",
        }
        for b in rows
    ]


@app.delete("/broadcasts/{broadcast_id}")
def delete_broadcast(
    broadcast_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role not in ("coordinator", "admin"):
        raise HTTPException(status_code=403, detail="Coordinator or admin access required.")
    broadcast = db.query(models.Broadcast).filter(models.Broadcast.id == broadcast_id).first()
    if not broadcast:
        raise HTTPException(status_code=404, detail="Broadcast not found.")
    db.delete(broadcast)
    db.commit()
    return {"message": "Broadcast deleted."}


@app.get("/groups/{group_id}/unread")
def get_group_unread(
    group_id: str,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")
    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")

    membership = db.query(models.GroupMember).filter(
        models.GroupMember.user_id == current_user.id,
        models.GroupMember.group_id == db_group.id
    ).first()

    # coordinators have no membership row — treat everything as read for them
    if not membership:
        return {"unread_messages": 0, "unread_tags": 0}

    last_read_id = membership.last_read_message_id

    if last_read_id is None:
        # never read anything — every message in this group is unread
        unread_messages = db.query(models.Message).filter(
            models.Message.group_id == db_group.id
        ).count()
    else:
        unread_messages = db.query(models.Message).filter(
            models.Message.group_id == db_group.id,
            models.Message.id > last_read_id
        ).count()

    unread_tags = db.query(models.Notification).filter(
        models.Notification.recipient_id == current_user.id,
        models.Notification.group_id == db_group.id,
        models.Notification.is_read == False
    ).count()

    return {"unread_messages": unread_messages, "unread_tags": unread_tags}


@app.post("/groups/{group_id}/read")
def mark_group_read(
    group_id: str,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")
    if not check_group_access(group_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied to this group")

    # find the latest message so we know what "read up to" means
    latest = db.query(models.Message).filter(
        models.Message.group_id == db_group.id
    ).order_by(models.Message.id.desc()).first()

    if not latest:
        return {"status": "ok"}

    membership = db.query(models.GroupMember).filter(
        models.GroupMember.user_id == current_user.id,
        models.GroupMember.group_id == db_group.id
    ).first()

    if membership:
        membership.last_read_message_id = latest.id
        db.commit()

    # also clear any tag notifications for this user in this group
    db.query(models.Notification).filter(
        models.Notification.recipient_id == current_user.id,
        models.Notification.group_id == db_group.id,
        models.Notification.is_read == False
    ).update({"is_read": True})
    db.commit()

    return {"status": "ok"}



@app.get("/coordinator/groups/overview")
def coordinator_groups_overview(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    One-stop overview of every group for the coordinator dashboard.
    Returns latest AI summary, latest student summary, and total human message count per group.
    Coordinator-only.
    """
    if current_user.role != "coordinator":
        raise HTTPException(status_code=403, detail="Coordinator access only")

    all_groups = db.query(models.Group).all()

    # reuse the same sort helper that /my-groups already uses
    def extract_group_number(name):
        match = re.search(r'\d+', name)
        return int(match.group()) if match else 999

    sorted_groups = sorted(all_groups, key=lambda g: extract_group_number(g.name))

    result = []
    for group in sorted_groups:
        # latest AI summary — group_id here is the string FK (e.g. "group-a")
        latest_ai = db.query(models.Summary).filter(
            models.Summary.group_id == group.string_id
        ).order_by(models.Summary.created_at.desc()).first()

        # latest student summary — also uses string_id FK
        latest_student = db.query(models.StudentSummary).filter(
            models.StudentSummary.group_id == group.string_id
        ).order_by(models.StudentSummary.created_at.desc()).first()

        # message count uses the integer PK — only human messages, no AI
        total_messages = db.query(models.Message).filter(
            models.Message.group_id == group.id,
            models.Message.is_AI == False
        ).count()

        result.append({
            "id": group.id,
            "name": group.name,
            "string_id": group.string_id,
            "ai_summary": {
                "summary_text": latest_ai.summary_text,
                "created_at": latest_ai.created_at.isoformat() + "Z",
            } if latest_ai else None,
            "student_summary": {
                "summary_text": latest_student.summary_text,
                "is_submitted": latest_student.is_submitted,
                "submitted_at": latest_student.submitted_at.isoformat() + "Z" if latest_student.submitted_at else None,
                "is_late": latest_student.is_late,
            } if latest_student else None,
            "total_messages": total_messages,
        })

    return result


@app.get("/coordinator/groups/{group_id}/contributions")
def coordinator_group_contributions(
    group_id: str,
    weeks: int = Query(4, ge=1, le=52, description="How many weeks back to look"),
    week_from: int | None = Query(None),
    week_to: int | None = Query(None),
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role != "coordinator":
        raise HTTPException(status_code=403, detail="Coordinator access only")

    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")

    if week_from is not None and week_to is not None:
        cp = db.query(models.CoursePeriod).order_by(models.CoursePeriod.created_at.desc()).first()
        if not cp:
            raise HTTPException(status_code=400, detail="No course period set.")
        start_d = cp.start_date + timedelta(days=(week_from - 1) * 7)
        end_d = cp.start_date + timedelta(days=week_to * 7)
        start_dt = datetime(start_d.year, start_d.month, start_d.day)
        end_dt = datetime(end_d.year, end_d.month, end_d.day)
    else:
        end_dt = datetime.utcnow()
        start_dt = end_dt - timedelta(weeks=weeks)

    # pull only human messages — ignore AI replies and system messages with no user
    messages = db.query(models.Message).filter(
        models.Message.group_id == db_group.id,
        models.Message.is_AI == False,
        models.Message.user_id != None,
        models.Message.timestamp >= start_dt,
        models.Message.timestamp < end_dt
    ).all()

    total_messages = len(messages)

    counts: dict[int, int] = {}
    for msg in messages:
        counts[msg.user_id] = counts.get(msg.user_id, 0) + 1

    contributions = []
    for user_id, count in counts.items():
        user = db.query(models.User).filter(models.User.id == user_id).first()
        username = user.username if user else f"user_{user_id}"
        percentage = round(count / total_messages * 100, 1) if total_messages > 0 else 0.0
        contributions.append({
            "username": username,
            "message_count": count,
            "percentage": percentage,
        })

    contributions.sort(key=lambda x: x["message_count"], reverse=True)

    return {
        "contributions": contributions,
        "total_messages": total_messages,
        "date_range": {
            "start": start_dt.isoformat() + "Z",
            "end": end_dt.isoformat() + "Z",
        },
        "week_from": week_from,
        "week_to": week_to,
    }


@app.get("/coordinator/groups/{group_id}/analysis")
def coordinator_group_analysis(
    group_id: str,
    weeks: int = Query(4, ge=1, le=52, description="How many weeks back to analyse"),
    week_from: int | None = Query(None),
    week_to: int | None = Query(None),
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role != "coordinator":
        raise HTTPException(status_code=403, detail="Coordinator access only")

    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")

    if week_from is not None and week_to is not None:
        cp = db.query(models.CoursePeriod).order_by(models.CoursePeriod.created_at.desc()).first()
        if not cp:
            raise HTTPException(status_code=400, detail="No course period set.")
        start_d = cp.start_date + timedelta(days=(week_from - 1) * 7)
        end_d = cp.start_date + timedelta(days=week_to * 7)
        start_dt = datetime(start_d.year, start_d.month, start_d.day)
        end_dt = datetime(end_d.year, end_d.month, end_d.day)
    else:
        end_dt = datetime.utcnow()
        start_dt = end_dt - timedelta(weeks=weeks)

    # format a UTC datetime as "27 Jun 2026, 23:59" in SGT (UTC+8)
    def _fmt_sgt(dt: datetime) -> str:
        return (dt + timedelta(hours=8)).strftime("%d %b %Y, %H:%M")

    # prefer a summary submitted within the selected period; fall back to most recent
    latest_student = db.query(models.StudentSummary).filter(
        models.StudentSummary.group_id == group_id,
        models.StudentSummary.created_at >= start_dt,
        models.StudentSummary.created_at <= end_dt,
    ).order_by(models.StudentSummary.created_at.desc()).first()

    if latest_student:
        summary_period_note = f"Submitted on {_fmt_sgt(latest_student.created_at)}"
    else:
        latest_student = db.query(models.StudentSummary).filter(
            models.StudentSummary.group_id == group_id
        ).order_by(models.StudentSummary.created_at.desc()).first()
        if latest_student:
            summary_period_note = (
                f"No summary was submitted for this period. "
                f"Using most recent submission from {_fmt_sgt(latest_student.created_at)}"
            )
        else:
            summary_period_note = "No student summary has been submitted yet."

    # true only when the student actually changed the AI draft before submitting
    edited_ai_copy = bool(
        latest_student
        and latest_student.summary_text
        and latest_student.ai_summary_copy
        and latest_student.summary_text != latest_student.ai_summary_copy
    )

    # all messages in the window, oldest first so the transcript reads naturally
    messages = db.query(models.Message).filter(
        models.Message.group_id == db_group.id,
        models.Message.timestamp >= start_dt,
        models.Message.timestamp < end_dt
    ).order_by(models.Message.timestamp.asc()).all()

    # no point calling GPT-4o if there's nothing to analyse
    if not messages:
        if week_from is not None and week_to is not None:
            no_msg_text = f"No messages found for Week {week_from} to Week {week_to}. Nothing to analyse."
        else:
            no_msg_text = "No messages found in the selected period. Nothing to analyse."
        return {
            "analysis_text": no_msg_text,
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "date_range": {
                "start": start_dt.isoformat() + "Z",
                "end": end_dt.isoformat() + "Z",
            },
            "week_from": week_from,
            "week_to": week_to,
            "summary_period_note": summary_period_note,
            "edited_ai_copy": edited_ai_copy,
            "student_summary_text": latest_student.summary_text if latest_student else None,
            "ai_summary_copy_text": latest_student.ai_summary_copy if latest_student else None,
        }

    # build a plain-text transcript — image messages get a placeholder
    transcript_lines = []
    for msg in messages:
        ts = msg.timestamp.strftime("%Y-%m-%d %H:%M")
        if msg.is_AI:
            sender = "AI Bot"
        else:
            sender = msg.user.username if msg.user else "Unknown"
        if (msg.message_type or "text") == "image":
            transcript_lines.append(f"[{ts}] {sender}: [sent an image]")
        else:
            transcript_lines.append(f"[{ts}] {sender}: {msg.content}")
    transcript = "\n".join(transcript_lines)

    if not openai_client:
        return {
            "analysis_text": "Error: OPENAI_API_KEY environment variable not set. Please configure your API key.",
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "date_range": {
                "start": start_dt.isoformat() + "Z",
                "end": end_dt.isoformat() + "Z",
            },
            "week_from": week_from,
            "week_to": week_to,
            "summary_period_note": summary_period_note,
            "edited_ai_copy": edited_ai_copy,
            "student_summary_text": latest_student.summary_text if latest_student else None,
            "ai_summary_copy_text": latest_student.ai_summary_copy if latest_student else None,
        }

    if latest_student:
        system_prompt = (
            "You are an academic coordinator reviewing a student group's progress. "
            "You have the group's chat transcript and their student-written weekly summary. "
            "Compare what actually happened in the chat with what the students reported. "
            "Be specific — cite actual topics, questions, or decisions from the chat when identifying gaps or discrepancies.\n\n"
            "Structure your response exactly like this (plain text, no markdown):\n\n"
            "What the students reported accurately:\n"
            "- (cite specific examples from the chat that match the summary)\n\n"
            "Gaps — things discussed in chat but missing from the summary:\n"
            "- (each gap labelled as Minor, Moderate, or Significant)\n\n"
            "Discrepancies — anything in the summary that does not match the chat:\n"
            "- (each discrepancy labelled as Minor, Moderate, or Significant, or 'None found' if everything checks out)\n\n"
            "Coordinator recommendations:\n"
            "- (2-3 specific actionable steps referencing actual topics from the chat)\n\n"
            f"Student-written summary:\n{latest_student.summary_text}"
        )
        if latest_student.ai_summary_copy:
            system_prompt += (
                f"\n\nAI-generated summary (the copy the students had access to when writing theirs):\n"
                f"{latest_student.ai_summary_copy}"
            )
    else:
        # no summary submitted — analyse the chat alone and flag the absence
        system_prompt = (
            "You are an academic coordinator reviewing a student group's progress. "
            "You have the group's chat transcript only — no student summary has been submitted. "
            "Analyse the chat and provide the following (plain text, no markdown):\n\n"
            "What the group has been working on:\n"
            "- (cite specific topics, tasks, or questions from the chat)\n\n"
            "Activity level:\n"
            "- Estimate how active the group is based on message frequency and depth of discussion "
            "(Very active / Moderately active / Low activity)\n"
            "- Note any periods of silence or sudden drop in participation\n\n"
            "Concerns visible from the chat:\n"
            "- (flag confusion, unresolved questions, dropped threads, or lack of progress — "
            "or 'None identified' if things look fine)\n\n"
            "Coordinator recommendations:\n"
            "- (2-3 specific actionable steps referencing actual topics from the chat)\n\n"
            "Note: no student summary has been submitted — flag this as a priority if the deadline is approaching."
        )

    user_prompt = f"Chat transcript ({weeks} week(s), {len(messages)} messages):\n\n{transcript}"
    if edited_ai_copy and latest_student and latest_student.ai_summary_copy:
        user_prompt += (
            "\n\nNote: the students edited the AI-generated summary before writing their own. "
            "Their edits may indicate what they chose to emphasise or omit."
        )

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            max_tokens=800,
            temperature=0.7
        )
        analysis_text = response.choices[0].message.content
    except Exception as e:
        analysis_text = f"Error: Failed to generate analysis. {str(e)}"

    return {
        "analysis_text": analysis_text,
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "date_range": {
            "start": start_dt.isoformat() + "Z",
            "end": end_dt.isoformat() + "Z",
        },
        "week_from": week_from,
        "week_to": week_to,
        "summary_period_note": summary_period_note,
        "edited_ai_copy": edited_ai_copy,
        "student_summary_text": latest_student.summary_text if latest_student else None,
        "ai_summary_copy_text": latest_student.ai_summary_copy if latest_student else None,
    }


@app.get("/supervisor/groups/overview")
def supervisor_groups_overview(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role != "supervisor":
        raise HTTPException(status_code=403, detail="Supervisor access only")

    memberships = db.query(models.GroupMember).filter(
        models.GroupMember.user_id == current_user.id
    ).all()
    member_group_ids = [m.group_id for m in memberships]
    groups = db.query(models.Group).filter(
        models.Group.id.in_(member_group_ids)
    ).all()

    def extract_group_number(name):
        match = re.search(r'\d+', name)
        return int(match.group()) if match else 999

    sorted_groups = sorted(groups, key=lambda g: extract_group_number(g.name))

    result = []
    for group in sorted_groups:
        latest_ai = db.query(models.Summary).filter(
            models.Summary.group_id == group.string_id
        ).order_by(models.Summary.created_at.desc()).first()

        latest_student = db.query(models.StudentSummary).filter(
            models.StudentSummary.group_id == group.string_id
        ).order_by(models.StudentSummary.created_at.desc()).first()

        total_messages = db.query(models.Message).filter(
            models.Message.group_id == group.id,
            models.Message.is_AI == False
        ).count()

        result.append({
            "id": group.id,
            "name": group.name,
            "string_id": group.string_id,
            "ai_summary": {
                "summary_text": latest_ai.summary_text,
                "created_at": latest_ai.created_at.isoformat() + "Z",
            } if latest_ai else None,
            "student_summary": {
                "summary_text": latest_student.summary_text,
                "is_submitted": latest_student.is_submitted,
                "submitted_at": latest_student.submitted_at.isoformat() + "Z" if latest_student.submitted_at else None,
                "is_late": latest_student.is_late,
            } if latest_student else None,
            "total_messages": total_messages,
        })

    return result


@app.get("/supervisor/groups/{group_id}/contributions")
def supervisor_group_contributions(
    group_id: str,
    weeks: int = Query(4, ge=1, le=52, description="How many weeks back to look"),
    week_from: int | None = Query(None),
    week_to: int | None = Query(None),
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role != "supervisor":
        raise HTTPException(status_code=403, detail="Supervisor access only")

    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")

    membership = db.query(models.GroupMember).filter(
        models.GroupMember.user_id == current_user.id,
        models.GroupMember.group_id == db_group.id
    ).first()
    if not membership:
        raise HTTPException(status_code=403, detail="Access denied")

    if week_from is not None and week_to is not None:
        cp = db.query(models.CoursePeriod).order_by(models.CoursePeriod.created_at.desc()).first()
        if not cp:
            raise HTTPException(status_code=400, detail="No course period set.")
        start_d = cp.start_date + timedelta(days=(week_from - 1) * 7)
        end_d = cp.start_date + timedelta(days=week_to * 7)
        start_dt = datetime(start_d.year, start_d.month, start_d.day)
        end_dt = datetime(end_d.year, end_d.month, end_d.day)
    else:
        end_dt = datetime.utcnow()
        start_dt = end_dt - timedelta(weeks=weeks)

    messages = db.query(models.Message).filter(
        models.Message.group_id == db_group.id,
        models.Message.is_AI == False,
        models.Message.user_id != None,
        models.Message.timestamp >= start_dt,
        models.Message.timestamp < end_dt
    ).all()

    total_messages = len(messages)

    counts: dict[int, int] = {}
    for msg in messages:
        counts[msg.user_id] = counts.get(msg.user_id, 0) + 1

    contributions = []
    for user_id, count in counts.items():
        user = db.query(models.User).filter(models.User.id == user_id).first()
        username = user.username if user else f"user_{user_id}"
        percentage = round(count / total_messages * 100, 1) if total_messages > 0 else 0.0
        contributions.append({
            "username": username,
            "message_count": count,
            "percentage": percentage,
        })

    contributions.sort(key=lambda x: x["message_count"], reverse=True)

    return {
        "contributions": contributions,
        "total_messages": total_messages,
        "date_range": {
            "start": start_dt.isoformat() + "Z",
            "end": end_dt.isoformat() + "Z",
        },
        "week_from": week_from,
        "week_to": week_to,
    }


@app.get("/supervisor/groups/{group_id}/analysis")
def supervisor_group_analysis(
    group_id: str,
    weeks: int = Query(4, ge=1, le=52),
    week_from: int | None = Query(None),
    week_to: int | None = Query(None),
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role != "supervisor":
        raise HTTPException(status_code=403, detail="Supervisor access only")

    db_group = db.query(models.Group).filter(models.Group.string_id == group_id).first()
    if not db_group:
        raise HTTPException(status_code=404, detail="Group not found")

    membership = db.query(models.GroupMember).filter(
        models.GroupMember.user_id == current_user.id,
        models.GroupMember.group_id == db_group.id
    ).first()
    if not membership:
        raise HTTPException(status_code=403, detail="Access denied")

    if week_from is not None and week_to is not None:
        cp = db.query(models.CoursePeriod).order_by(models.CoursePeriod.created_at.desc()).first()
        if not cp:
            raise HTTPException(status_code=400, detail="No course period set.")
        start_d = cp.start_date + timedelta(days=(week_from - 1) * 7)
        end_d = cp.start_date + timedelta(days=week_to * 7)
        start_dt = datetime(start_d.year, start_d.month, start_d.day)
        end_dt = datetime(end_d.year, end_d.month, end_d.day)
    else:
        end_dt = datetime.utcnow()
        start_dt = end_dt - timedelta(weeks=weeks)

    def _fmt_sgt(dt: datetime) -> str:
        return (dt + timedelta(hours=8)).strftime("%d %b %Y, %H:%M")

    latest_student = db.query(models.StudentSummary).filter(
        models.StudentSummary.group_id == group_id,
        models.StudentSummary.created_at >= start_dt,
        models.StudentSummary.created_at <= end_dt,
    ).order_by(models.StudentSummary.created_at.desc()).first()

    if latest_student:
        summary_period_note = f"Submitted on {_fmt_sgt(latest_student.created_at)}"
    else:
        latest_student = db.query(models.StudentSummary).filter(
            models.StudentSummary.group_id == group_id
        ).order_by(models.StudentSummary.created_at.desc()).first()
        if latest_student:
            summary_period_note = (
                f"No summary was submitted for this period. "
                f"Using most recent submission from {_fmt_sgt(latest_student.created_at)}"
            )
        else:
            summary_period_note = "No student summary has been submitted yet."

    edited_ai_copy = bool(
        latest_student
        and latest_student.summary_text
        and latest_student.ai_summary_copy
        and latest_student.summary_text != latest_student.ai_summary_copy
    )

    messages = db.query(models.Message).filter(
        models.Message.group_id == db_group.id,
        models.Message.timestamp >= start_dt,
        models.Message.timestamp < end_dt
    ).order_by(models.Message.timestamp.asc()).all()

    if not messages:
        if week_from is not None and week_to is not None:
            no_msg_text = f"No messages found for Week {week_from} to Week {week_to}. Nothing to analyse."
        else:
            no_msg_text = "No messages found in the selected period. Nothing to analyse."
        return {
            "analysis_text": no_msg_text,
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "date_range": {
                "start": start_dt.isoformat() + "Z",
                "end": end_dt.isoformat() + "Z",
            },
            "week_from": week_from,
            "week_to": week_to,
            "summary_period_note": summary_period_note,
            "edited_ai_copy": edited_ai_copy,
            "student_summary_text": latest_student.summary_text if latest_student else None,
            "ai_summary_copy_text": latest_student.ai_summary_copy if latest_student else None,
        }

    transcript_lines = []
    for msg in messages:
        ts = msg.timestamp.strftime("%Y-%m-%d %H:%M")
        if msg.is_AI:
            sender = "AI Bot"
        else:
            sender = msg.user.username if msg.user else "Unknown"
        if (msg.message_type or "text") == "image":
            transcript_lines.append(f"[{ts}] {sender}: [sent an image]")
        else:
            transcript_lines.append(f"[{ts}] {sender}: {msg.content}")
    transcript = "\n".join(transcript_lines)

    if not openai_client:
        return {
            "analysis_text": "Error: OPENAI_API_KEY environment variable not set. Please configure your API key.",
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "date_range": {
                "start": start_dt.isoformat() + "Z",
                "end": end_dt.isoformat() + "Z",
            },
            "week_from": week_from,
            "week_to": week_to,
            "summary_period_note": summary_period_note,
            "edited_ai_copy": edited_ai_copy,
            "student_summary_text": latest_student.summary_text if latest_student else None,
            "ai_summary_copy_text": latest_student.ai_summary_copy if latest_student else None,
        }

    if latest_student:
        system_prompt = (
            "You are an academic supervisor reviewing a student group's progress. "
            "You have the group's chat transcript and their student-written weekly summary. "
            "Compare what actually happened in the chat with what the students reported. "
            "Be specific — cite actual topics, questions, or decisions from the chat when identifying gaps or discrepancies.\n\n"
            "Structure your response exactly like this (plain text, no markdown):\n\n"
            "What the students reported accurately:\n"
            "- (cite specific examples from the chat that match the summary)\n\n"
            "Gaps — things discussed in chat but missing from the summary:\n"
            "- (each gap labelled as Minor, Moderate, or Significant)\n\n"
            "Discrepancies — anything in the summary that does not match the chat:\n"
            "- (each discrepancy labelled as Minor, Moderate, or Significant, or 'None found' if everything checks out)\n\n"
            "Coordinator recommendations:\n"
            "- (2-3 specific actionable steps referencing actual topics from the chat)\n\n"
            f"Student-written summary:\n{latest_student.summary_text}"
        )
        if latest_student.ai_summary_copy:
            system_prompt += (
                f"\n\nAI-generated summary (the copy the students had access to when writing theirs):\n"
                f"{latest_student.ai_summary_copy}"
            )
    else:
        system_prompt = (
            "You are an academic supervisor reviewing a student group's progress. "
            "You have the group's chat transcript only — no student summary has been submitted. "
            "Analyse the chat and provide the following (plain text, no markdown):\n\n"
            "What the group has been working on:\n"
            "- (cite specific topics, tasks, or questions from the chat)\n\n"
            "Activity level:\n"
            "- Estimate how active the group is based on message frequency and depth of discussion "
            "(Very active / Moderately active / Low activity)\n"
            "- Note any periods of silence or sudden drop in participation\n\n"
            "Concerns visible from the chat:\n"
            "- (flag confusion, unresolved questions, dropped threads, or lack of progress — "
            "or 'None identified' if things look fine)\n\n"
            "Coordinator recommendations:\n"
            "- (2-3 specific actionable steps referencing actual topics from the chat)\n\n"
            "Note: no student summary has been submitted — flag this as a priority if the deadline is approaching."
        )

    user_prompt = f"Chat transcript ({weeks} week(s), {len(messages)} messages):\n\n{transcript}"
    if edited_ai_copy and latest_student and latest_student.ai_summary_copy:
        user_prompt += (
            "\n\nNote: the students edited the AI-generated summary before writing their own. "
            "Their edits may indicate what they chose to emphasise or omit."
        )

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            max_tokens=800,
            temperature=0.7
        )
        analysis_text = response.choices[0].message.content
    except Exception as e:
        analysis_text = f"Error: Failed to generate analysis. {str(e)}"

    return {
        "analysis_text": analysis_text,
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "date_range": {
            "start": start_dt.isoformat() + "Z",
            "end": end_dt.isoformat() + "Z",
        },
        "week_from": week_from,
        "week_to": week_to,
        "summary_period_note": summary_period_note,
        "edited_ai_copy": edited_ai_copy,
        "student_summary_text": latest_student.summary_text if latest_student else None,
        "ai_summary_copy_text": latest_student.ai_summary_copy if latest_student else None,
    }


VALID_ROLES = {"student", "supervisor", "coordinator", "admin"}


def _user_dict(user: models.User) -> dict:
    # shared shape returned by all admin user endpoints
    return {
        "id": user.id,
        "username": user.username,
        "email": user.email,
        "full_name": user.full_name,
        "role": user.role,
        "is_active": user.is_active,
        "created_at": user.created_at.isoformat() + "Z",
    }


@app.get("/admin/users")
def admin_list_users(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    users = db.query(models.User).order_by(models.User.created_at.asc()).all()
    return [_user_dict(u) for u in users]


@app.post("/admin/users", status_code=201)
def admin_create_user(
    body: CreateUserRequest,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    if body.role not in VALID_ROLES:
        raise HTTPException(status_code=400, detail=f"Invalid role. Must be one of: {sorted(VALID_ROLES)}")
    if db.query(models.User).filter(models.User.username == body.username).first():
        raise HTTPException(status_code=400, detail="Username already taken.")
    if body.email and db.query(models.User).filter(models.User.email == body.email).first():
        raise HTTPException(status_code=400, detail="Email already taken.")
    user = models.User(
        username=body.username,
        password_hash=hash_password(body.password),
        role=body.role,
        email=body.email,
        full_name=body.full_name,
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return _user_dict(user)


@app.put("/admin/users/{user_id}")
def admin_update_user(
    user_id: int,
    body: UpdateUserRequest,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found.")
    if body.role is not None:
        if body.role not in VALID_ROLES:
            raise HTTPException(status_code=400, detail=f"Invalid role. Must be one of: {sorted(VALID_ROLES)}")
        user.role = body.role
    if body.email is not None:
        existing = db.query(models.User).filter(
            models.User.email == body.email,
            models.User.id != user_id,
        ).first()
        if existing:
            raise HTTPException(status_code=400, detail="Email already taken.")
        user.email = body.email
    if body.full_name is not None:
        user.full_name = body.full_name
    if body.is_active is not None:
        user.is_active = body.is_active
    db.commit()
    db.refresh(user)
    return _user_dict(user)


@app.delete("/admin/users/{user_id}")
def admin_deactivate_user(
    user_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="Cannot deactivate your own account.")
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found.")
    user.is_active = False
    db.commit()
    return {"message": "User deactivated."}


@app.delete("/admin/users/{user_id}/permanent")
def admin_hard_delete_user(
    user_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="Cannot delete your own account.")
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found.")
    # notifications hold non-nullable FKs to both users and messages — must go first
    db.query(models.Notification).filter(
        (models.Notification.recipient_id == user_id) | (models.Notification.sender_id == user_id)
    ).delete(synchronize_session=False)
    db.query(models.GroupMember).filter(models.GroupMember.user_id == user_id).delete()
    db.query(models.PushToken).filter(models.PushToken.user_id == user_id).delete()
    db.delete(user)
    db.commit()
    return {"message": "User permanently deleted."}


@app.post("/admin/users/bulk", status_code=201)
def admin_bulk_create_users(
    body: list[BulkUserEntry],
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    created = []
    skipped = []
    for entry in body:
        if db.query(models.User).filter(models.User.username == entry.username).first():
            skipped.append(entry.username)
            continue
        if entry.role not in VALID_ROLES:
            skipped.append(entry.username)
            continue
        user = models.User(
            username=entry.username,
            password_hash=hash_password(entry.password),
            role=entry.role,
            email=entry.email,
            full_name=entry.full_name,
        )
        db.add(user)
        db.flush()  # get user.id before creating the membership below
        if entry.group_id is not None:
            group = db.query(models.Group).filter(models.Group.id == entry.group_id).first()
            if group:
                db.add(models.GroupMember(user_id=user.id, group_id=entry.group_id))
        created.append(entry.username)
    db.commit()
    return {"created": created, "skipped": skipped}


@app.post("/admin/groups/{group_id}/members", status_code=201)
def admin_add_group_member(
    group_id: int,
    body: AddGroupMemberRequest,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    group = db.query(models.Group).filter(models.Group.id == group_id).first()
    if not group:
        raise HTTPException(status_code=404, detail="Group not found.")
    user = db.query(models.User).filter(models.User.id == body.user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found.")
    already = db.query(models.GroupMember).filter(
        models.GroupMember.user_id == body.user_id,
        models.GroupMember.group_id == group_id,
    ).first()
    if already:
        return {"message": "User already in group."}
    db.add(models.GroupMember(
        user_id=body.user_id,
        group_id=group_id,
        role_in_group=body.role_in_group,
    ))
    db.commit()
    return {"message": "User added to group."}


@app.delete("/admin/groups/{group_id}/members/{user_id}")
def admin_remove_group_member(
    group_id: int,
    user_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    membership = db.query(models.GroupMember).filter(
        models.GroupMember.user_id == user_id,
        models.GroupMember.group_id == group_id,
    ).first()
    if not membership:
        raise HTTPException(status_code=404, detail="Membership not found.")
    db.delete(membership)
    db.commit()
    return {"message": "User removed from group."}


@app.get("/admin/groups")
def admin_list_groups(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    all_groups = db.query(models.Group).all()
    # same numeric sort used by /my-groups for coordinators
    def extract_group_number(group_name):
        match = re.search(r"\d+", group_name)
        return int(match.group()) if match else 999
    sorted_groups = sorted(all_groups, key=lambda g: extract_group_number(g.name))
    return [{"id": g.id, "name": g.name, "string_id": g.string_id} for g in sorted_groups]


@app.post("/admin/groups", status_code=201)
def admin_create_group(
    body: CreateGroupRequest,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    name = body.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="Group name cannot be empty.")
    if len(name) > 100:
        raise HTTPException(status_code=400, detail="Group name must be 100 characters or fewer.")
    string_id = re.sub(r"[^a-z0-9-]", "", name.lower().replace(" ", "-"))
    if db.query(models.Group).filter(models.Group.string_id == string_id).first():
        raise HTTPException(status_code=400, detail="A group with this name already exists.")
    group = models.Group(name=name, string_id=string_id)
    db.add(group)
    db.commit()
    db.refresh(group)
    return {"id": group.id, "name": group.name, "string_id": group.string_id}


@app.delete("/admin/groups/{group_id}")
def admin_delete_group(
    group_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    group = db.query(models.Group).filter(models.Group.id == group_id).first()
    if not group:
        raise HTTPException(status_code=404, detail="Group not found.")

    # notifications first — they hold a FK to messages.id, so they must go before messages
    db.query(models.Notification).filter(models.Notification.group_id == group_id).delete()
    db.query(models.GroupMember).filter(models.GroupMember.group_id == group_id).delete()
    db.query(models.Message).filter(models.Message.group_id == group_id).delete()
    # these three use string_id as their FK, not the integer id
    db.query(models.Document).filter(models.Document.group_id == group.string_id).delete()
    db.query(models.Summary).filter(models.Summary.group_id == group.string_id).delete()
    db.query(models.StudentSummary).filter(models.StudentSummary.group_id == group.string_id).delete()
    db.delete(group)
    db.commit()
    return {"message": "Group and all associated data permanently deleted."}


@app.get("/admin/groups/{group_id}/members")
def admin_list_group_members(
    group_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    group = db.query(models.Group).filter(models.Group.id == group_id).first()
    if not group:
        raise HTTPException(status_code=404, detail="Group not found.")
    members = (
        db.query(models.User)
        .join(models.GroupMember, models.GroupMember.user_id == models.User.id)
        .filter(models.GroupMember.group_id == group_id)
        .all()
    )
    return [
        {
            "id": m.id,
            "username": m.username,
            "full_name": m.full_name,
            "role": m.role,
        }
        for m in members
    ]


@app.post("/admin/students/import", status_code=201)
def admin_import_students(
    body: list[ExcelStudentEntry],
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    created = []
    skipped = []
    errors = []
    for entry in body:
        username = entry.username.lower().strip()
        email = f"{username}@e.ntu.edu.sg"
        group_name = f"Group {entry.group_id}"
        existing = db.query(models.User).filter(models.User.username == username).first()
        if existing:
            if existing.is_active:
                skipped.append(username)
                continue
            # reactivate an inactive account and refresh their details
            existing.is_active = True
            existing.full_name = entry.full_name
            existing.email = email
            if entry.student_id:
                existing.student_id = entry.student_id
            user = existing
        else:
            # Skip if email is already taken by a different active account
            if db.query(models.User).filter(models.User.email == email).first():
                skipped.append(username)
                continue
            user = models.User(
                username=username,
                password_hash=hash_password(str(uuid.uuid4())),
                role="student",
                email=email,
                full_name=entry.full_name,
                student_id=entry.student_id if entry.student_id else None,
            )
            db.add(user)
            db.flush()
        # case-insensitive group lookup — create the group if it doesn't exist yet
        group = db.query(models.Group).filter(
            models.Group.name.ilike(group_name)
        ).first()
        if not group:
            string_id = group_name.lower().replace(" ", "-")
            group = models.Group(name=group_name, string_id=string_id)
            db.add(group)
            db.flush()
        # link student to group
        already = db.query(models.GroupMember).filter(
            models.GroupMember.user_id == user.id,
            models.GroupMember.group_id == group.id,
        ).first()
        if not already:
            db.add(models.GroupMember(user_id=user.id, group_id=group.id))
        created.append(username)
    db.commit()
    return {"created": created, "skipped": skipped, "errors": errors}


@app.post("/feedback", status_code=201)
def submit_feedback(
    body: FeedbackRequest,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if not body.content.strip():
        raise HTTPException(status_code=400, detail="Content cannot be empty.")
    if len(body.content) > 2000:
        raise HTTPException(status_code=400, detail="Content must be 2000 characters or fewer.")
    if body.feedback_type not in ("general", "bug"):
        raise HTTPException(status_code=400, detail="feedback_type must be 'general' or 'bug'.")
    item = models.Feedback(
        content=body.content,
        feedback_type=body.feedback_type,
        submitted_by_user_id=current_user.id,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return {
        "id": item.id,
        "content": item.content,
        "feedback_type": item.feedback_type,
        "submitted_by": current_user.username,
        "created_at": item.created_at.isoformat() + "Z",
    }


@app.get("/admin/feedback")
def get_all_feedback(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    rows = (
        db.query(models.Feedback)
        .order_by(models.Feedback.created_at.desc())
        .all()
    )
    return [
        {
            "id": f.id,
            "content": f.content,
            "feedback_type": f.feedback_type,
            "submitted_by": f.submitted_by.username if f.submitted_by else None,
            "is_resolved": f.is_resolved,
            "created_at": f.created_at.isoformat() + "Z",
        }
        for f in rows
    ]


@app.put("/admin/feedback/{feedback_id}/resolve")
def toggle_feedback_resolved(
    feedback_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    item = db.query(models.Feedback).filter(models.Feedback.id == feedback_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Feedback not found.")
    item.is_resolved = not item.is_resolved
    db.commit()
    db.refresh(item)
    return {
        "id": item.id,
        "content": item.content,
        "feedback_type": item.feedback_type,
        "submitted_by": item.submitted_by.username if item.submitted_by else None,
        "is_resolved": item.is_resolved,
        "created_at": item.created_at.isoformat() + "Z",
    }
